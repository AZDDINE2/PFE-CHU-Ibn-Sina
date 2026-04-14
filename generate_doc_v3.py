# -*- coding: utf-8 -*-
"""
Documentation complète du projet PFE — CHU Ibn Sina
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Marges ──────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ─────────────────────────────────────────────────────────
def shade(cell, hex_color):
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear')
    s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(s)

def tbl_grid(rows, cols):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'
    for row in t.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t

def header_row(tbl, labels, bg='1F497D'):
    for i, label in enumerate(labels):
        cell = tbl.rows[0].cells[i]
        cell.text = label
        shade(cell, bg)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

def h1(text):
    p = doc.add_heading(text, 1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 73, 125)
    return p

def h2(text):
    p = doc.add_heading(text, 2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(68, 114, 196)
    return p

def h3(text):
    return doc.add_heading(text, 3)

def para(text, bold=False, italic=False, size=11, color=None, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)

def code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Cm(1)
    return p

def fill_table(tbl, data, start_row=1):
    for i, row_data in enumerate(data):
        row = tbl.rows[start_row + i]
        for j, val in enumerate(row_data):
            row.cells[j].text = str(val)
            for p in row.cells[j].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

def alt_rows(tbl, start=1, color='EBF3FB'):
    for i in range(start, len(tbl.rows), 2):
        for cell in tbl.rows[i].cells:
            shade(cell, color)

# ════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════
doc.add_paragraph('\n\n')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PROJET DE FIN D\'ETUDES')
r.bold = True; r.font.size = Pt(22)
r.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Tableau de Bord Decisionnel')
r2.bold = True; r2.font.size = Pt(18)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('CHU Ibn Sina -- Urgences Hospitalieres')
r3.bold = True; r3.font.size = Pt(16)
r3.font.color.rgb = RGBColor(192, 0, 0)

doc.add_paragraph('\n\n')

info = [
    ('Realise par',      'Azddine'),
    ('Encadrant',        '---'),
    ('Etablissement',    'Ecole d\'Ingenieurs'),
    ('Annee academique', '2024 / 2025'),
    ('Date',             datetime.date.today().strftime('%d/%m/%Y')),
    ('Version',          '1.0'),
]
tbl_garde = tbl_grid(len(info), 2)
for i, (k, v) in enumerate(info):
    tbl_garde.rows[i].cells[0].text = k
    tbl_garde.rows[i].cells[1].text = v
    tbl_garde.rows[i].cells[0].width = Cm(6)
    tbl_garde.rows[i].cells[1].width = Cm(8)
    for p in tbl_garde.rows[i].cells[0].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(31, 73, 125)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# RÉSUMÉ EXÉCUTIF
# ════════════════════════════════════════════════════════════════════
h1('Resume Executif')
para(
    'Ce projet consiste en la conception et le developpement d\'un tableau de bord decisionnel '
    'complet pour le Centre Hospitalier Universitaire (CHU) Ibn Sina de Rabat. '
    'L\'application permet aux responsables medicaux et administratifs de visualiser, analyser '
    'et predire les flux de patients aux urgences en temps reel.'
)
doc.add_paragraph()
para('Chiffres cles :', bold=True)
tbl_resume = tbl_grid(6, 2)
header_row(tbl_resume, ['Indicateur', 'Valeur'])
fill_table(tbl_resume, [
    ('Visites aux urgences traitees',   '530 898'),
    ('Actes medicaux (soins)',          '737 512'),
    ('Etablissements hospitaliers',     '8 hopitaux'),
    ('Periode couverte',                '2019 -- 2026 (7 ans)'),
    ('Meilleur score R2 (XGBoost)',     '96.70%'),
])
alt_rows(tbl_resume)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ════════════════════════════════════════════════════════════════════
h1('Table des Matieres')
toc = [
    ('1', 'Contexte et Objectifs'),
    ('2', 'Architecture Technique'),
    ('3', 'Dataset'),
    ('4', 'Structure de la Base de Donnees'),
    ('5', 'Relations entre les Tables'),
    ('6', 'Pipeline ETL'),
    ('7', 'Backend -- API FastAPI'),
    ('8', 'Frontend -- Dashboard React'),
    ('9', 'Modeles Machine Learning'),
    ('10', 'Fonctionnalites du Dashboard'),
    ('11', 'Securite et Authentification'),
    ('12', 'Deploiement Docker'),
    ('13', 'Performances et Resultats'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'{num}.  {title}')
    r.font.size = Pt(11)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 1. CONTEXTE ET OBJECTIFS
# ════════════════════════════════════════════════════════════════════
h1('1. Contexte et Objectifs')
para(
    'Le CHU Ibn Sina de Rabat est le plus grand hopital universitaire du Maroc avec plus de '
    '924 lits, 520 medecins et 850 infirmiers. Ses services d\'urgences traitent des dizaines '
    'de milliers de patients par an. La gestion de ces flux massifs necessite des outils '
    'analytiques avances pour optimiser les ressources et ameliorer la qualite des soins.'
)

h2('Objectifs du projet')
objectives = [
    'Centraliser les donnees de 8 etablissements hospitaliers dans une base unifiee',
    'Visualiser les KPIs cles : flux patients, taux d\'hospitalisation, duree de sejour',
    'Analyser les patterns temporels (heures de pic, saisonnalite, tendances)',
    'Predire l\'affluence future avec Prophet (R2 = 95.57%)',
    'Estimer la duree de sejour avec Random Forest et XGBoost (R2 = 96.70%)',
    'Planifier les ressources humaines et materielles pour une date donnee',
    'Fournir une interface responsive avec mode sombre, export PDF et envoi email',
    'Securiser l\'acces avec JWT et un systeme de roles (admin, medecin, infirmier, direction)',
]
for obj in objectives:
    bullet(obj)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE TECHNIQUE
# ════════════════════════════════════════════════════════════════════
h1('2. Architecture Technique')
para('L\'application suit une architecture client-serveur containerisee :')
doc.add_paragraph()

h2('Stack technologique')
tbl_stack = tbl_grid(7, 3)
header_row(tbl_stack, ['Couche', 'Technologie', 'Role'])
fill_table(tbl_stack, [
    ('Interface utilisateur', 'React 18 + TypeScript',         'Dashboard interactif, graphiques, formulaires'),
    ('Serveur web',           'Nginx Alpine',                   'Serveur fichiers statiques + proxy API'),
    ('API REST',              'Python 3.11 + FastAPI',          'Logique metier, endpoints, ML'),
    ('Base de donnees',       'SQLite 3 (mode WAL)',            'Stockage persistant des donnees'),
    ('Machine Learning',      'Prophet, Scikit-learn, XGBoost', 'Previsions et predictions'),
    ('Containerisation',      'Docker + Docker Compose',        'Deploiement portable et reproductible'),
])
alt_rows(tbl_stack)

doc.add_paragraph()
h2('Flux de donnees')
code(
    'Navigateur (port 3000)\n'
    '    --> Nginx (proxy)\n'
    '        --> FastAPI (port 8000)\n'
    '            --> SQLite (chu.db)\n'
    '            --> Modeles ML (.pkl)\n'
    '                <-- Reponse JSON\n'
    '    <-- Rendu React'
)

doc.add_paragraph()
h2('Conteneurs Docker')
tbl_docker = tbl_grid(3, 4)
header_row(tbl_docker, ['Conteneur', 'Image de base', 'Port expose', 'Volumes montes'])
fill_table(tbl_docker, [
    ('chu_backend',  'python:3.11-slim', '8000', './data:/app/data   ./models:/app/models'),
    ('chu_frontend', 'nginx:alpine',     '3000', 'aucun (build statique)'),
])
alt_rows(tbl_docker)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 3. DATASET
# ════════════════════════════════════════════════════════════════════
h1('3. Dataset')
para(
    'Le dataset est un jeu de donnees simule reproduisant fidelement le comportement '
    'des urgences hospitalieres marocaines. Les distributions statistiques (ages, '
    'triages, orientations, saisonnalite) respectent les rapports publics du '
    'Ministere de la Sante du Maroc.'
)

h2('Volume global')
tbl_vol = tbl_grid(6, 4)
header_row(tbl_vol, ['Table', 'Lignes', 'Colonnes', 'Description'])
fill_table(tbl_vol, [
    ('urgences',       '530 898', '31', 'Passages aux urgences (2019-2026)'),
    ('soins',          '737 512', '13', 'Actes medicaux associes aux passages'),
    ('etablissements', '8',       '24', 'Hopitaux avec KPIs calcules'),
    ('personnel',      '1 086',   '12', 'Medecins et infirmiers'),
    ('lits',           '3 295',   '10', 'Lits hospitaliers et occupation'),
])
alt_rows(tbl_vol)

doc.add_paragraph()
h2('Colonnes de la table urgences (31 colonnes)')

h3('Identification patient')
tbl_id = tbl_grid(9, 3)
header_row(tbl_id, ['Colonne', 'Type', 'Description'], '2E75B6')
fill_table(tbl_id, [
    ('IPP',            'FLOAT',   'Identifiant Permanent du Patient'),
    ('Nom_Complet',    'TEXT',    'Nom et prenom du patient'),
    ('Age',            'INTEGER', 'Age en annees (0-110)'),
    ('Sexe',           'TEXT',    'M ou F'),
    ('CIN',            'TEXT',    'Carte d\'identite nationale'),
    ('Groupe_Sanguin', 'TEXT',    'A+, A-, B+, B-, O+, O-, AB+, AB-'),
    ('Antecedents',    'TEXT',    'Antecedents medicaux du patient'),
    ('Mutuelle',       'TEXT',    'Couverture d\'assurance maladie'),
])
alt_rows(tbl_id)

doc.add_paragraph()
h3('Etablissement hospitalier')
tbl_etab = tbl_grid(4, 3)
header_row(tbl_etab, ['Colonne', 'Type', 'Description'], '2E75B6')
fill_table(tbl_etab, [
    ('Etablissement', 'TEXT', 'Nom de l\'hopital'),
    ('Type_Etab',     'TEXT', 'CHU / Hopital regional / Centre de sante'),
    ('Ville',         'TEXT', 'Ville de l\'etablissement'),
])
alt_rows(tbl_etab)

doc.add_paragraph()
h3('Passage aux urgences')
tbl_pass = tbl_grid(11, 3)
header_row(tbl_pass, ['Colonne', 'Type', 'Description'], '2E75B6')
fill_table(tbl_pass, [
    ('Date_Arrivee',       'DATETIME', 'Date et heure d\'arrivee aux urgences'),
    ('Date_Sortie',        'DATETIME', 'Date et heure de sortie'),
    ('Niveau_Triage',      'TEXT',     'P1 Critique / P2 Urgent / P3 Semi-urgent / P4 Non urgent'),
    ('Motif_Consultation', 'TEXT',     'Motif de la visite (douleur, accident, fievre...)'),
    ('Orientation',        'TEXT',     'Hospitalisation / Retour domicile / Deces / Fugue...'),
    ('Duree_Sejour_min',   'INTEGER',  'Duree du passage en minutes (0 a 4320)'),
    ('Nb_Medecins_Dispo',  'INTEGER',  'Nombre de medecins disponibles ce jour'),
    ('Nb_Lits_Dispo',      'INTEGER',  'Nombre de lits disponibles ce jour'),
    ('Prix_Sejour',        'INTEGER',  'Cout du sejour en dirhams (MAD)'),
    ('Prix_Soins',         'INTEGER',  'Cout des soins en dirhams (MAD)'),
])
alt_rows(tbl_pass)

doc.add_paragraph()
h3('Features calculees par l\'ETL (Feature Engineering)')
tbl_feat = tbl_grid(11, 3)
header_row(tbl_feat, ['Colonne', 'Type', 'Calcul / Description'], '375623')
fill_table(tbl_feat, [
    ('Heure_Arrivee',  'FLOAT',   'date_arrivee.dt.hour  -- heure (0-23)'),
    ('Jour_Semaine',   'FLOAT',   'date_arrivee.dt.dayofweek  -- 0=Lundi, 6=Dimanche'),
    ('Mois',           'FLOAT',   'date_arrivee.dt.month  -- 1 a 12'),
    ('Annee',          'FLOAT',   'date_arrivee.dt.year  -- 2019 a 2026'),
    ('Saison',         'TEXT',    'Regle sur le mois : Hiver/Printemps/Ete/Automne'),
    ('Tranche_Horaire','TEXT',    'Regle heure : Matin/Apres-midi/Soir/Nuit'),
    ('Nom_Jour',       'TEXT',    'Lundi, Mardi, Mercredi...'),
    ('Groupe_Age',     'TEXT',    'Regle age : Enfant/Ado/Adulte Jeune/Adulte/Senior'),
    ('Jour_Ferie',     'INTEGER', '1 si jour ferie marocain, 0 sinon'),
    ('Est_Pic',        'INTEGER', '1 si flux horaire > Q75 (75e percentile), 0 sinon'),
])
alt_rows(tbl_feat)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 4. STRUCTURE DE LA BASE DE DONNÉES
# ════════════════════════════════════════════════════════════════════
h1('4. Structure de la Base de Donnees')
para(
    'La base de donnees chu.db utilise SQLite en mode WAL (Write-Ahead Logging) '
    'pour permettre des lectures simultanees sans blocage. Taille totale : ~350 MB.'
)
code(
    'PRAGMA journal_mode = WAL\n'
    'PRAGMA synchronous  = NORMAL\n'
    'PRAGMA busy_timeout = 120000'
)

doc.add_paragraph()
h2('Colonnes de la table soins (13 colonnes)')
tbl_soins = tbl_grid(14, 3)
header_row(tbl_soins, ['Colonne', 'Type', 'Description'])
fill_table(tbl_soins, [
    ('id_soin',       'INTEGER', 'Identifiant unique de l\'acte medical (PK)'),
    ('IPP',           'INTEGER', 'Reference vers urgences.IPP (FK)'),
    ('Etablissement', 'TEXT',    'Hopital ou l\'acte a ete realise'),
    ('Age',           'INTEGER', 'Age du patient'),
    ('Sexe',          'TEXT',    'M ou F'),
    ('niveau_triage', 'TEXT',    'P1/P2/P3/P4'),
    ('Type_Soin',     'TEXT',    'Consultation / Chirurgie / Biologie / Radiologie...'),
    ('Medicament',    'TEXT',    'Medicament administre'),
    ('Medecin',       'TEXT',    'Medecin responsable'),
    ('Cout_Soin',     'FLOAT',   'Cout de l\'acte en dirhams (MAD)'),
    ('Duree_Soin_Min','INTEGER', 'Duree de l\'acte en minutes'),
    ('Resultat',      'TEXT',    'Stabilise / Gueri / Transfere / Decede...'),
    ('Date_Soin',     'TEXT',    'Date de realisation de l\'acte'),
])
alt_rows(tbl_soins)

doc.add_paragraph()
h2('Colonnes de la table etablissements (24 colonnes)')
tbl_etabs = tbl_grid(13, 3)
header_row(tbl_etabs, ['Colonne', 'Type', 'Description'])
fill_table(tbl_etabs, [
    ('nom',                 'TEXT',    'Nom de l\'hopital (PK)'),
    ('type_etab',           'TEXT',    'Type d\'etablissement'),
    ('ville',               'TEXT',    'Ville'),
    ('capacite_lits',       'INTEGER', 'Capacite totale en lits'),
    ('nb_medecins',         'INTEGER', 'Nombre de medecins'),
    ('nb_urgentistes',      'INTEGER', 'Nombre d\'urgentistes'),
    ('Nb_Patients',         'INTEGER', 'Total passages (calcule depuis urgences)'),
    ('Duree_Moy_Min',       'FLOAT',   'Duree moyenne de sejour en minutes'),
    ('Taux_Hospit_Pct',     'FLOAT',   '% patients hospitalises'),
    ('Taux_Fugue_Pct',      'FLOAT',   '% patients partis sans soins'),
    ('Taux_P1_Pct',         'FLOAT',   '% cas critiques (P1)'),
    ('Alerte_Charge',       'TEXT',    'Critique / Eleve / Normal (regle sur KPIs)'),
])
alt_rows(tbl_etabs)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 5. RELATIONS ENTRE LES TABLES
# ════════════════════════════════════════════════════════════════════
h1('5. Relations entre les Tables')

h2('Schema relationnel')
code(
    'urgences                        soins\n'
    '-------------------             -------------------\n'
    'IPP (PK)        <-----------    IPP (FK)\n'
    'Etablissement   <----+          id_soin (PK)\n'
    'Nom_Complet          |          Etablissement (FK) --+\n'
    'Age                  |          Type_Soin            |\n'
    'Niveau_Triage        |          Cout_Soin            |\n'
    'Orientation          |          Resultat             |\n'
    '...                  |                               |\n'
    '                     |    etablissements             |\n'
    '                     |    -------------------        |\n'
    '                     +--- nom (PK)        <----------+\n'
    '                          type_etab\n'
    '                          capacite_lits\n'
    '                          nb_medecins\n'
    '                          Alerte_Charge'
)

doc.add_paragraph()
h2('Description des relations')
tbl_rel = tbl_grid(5, 4)
header_row(tbl_rel, ['Relation', 'Cle source', 'Cle cible', 'Cardinalite'])
fill_table(tbl_rel, [
    ('urgences -> soins',
     'urgences.IPP',
     'soins.IPP',
     '1 -- N  (1 passage genere plusieurs actes medicaux)'),
    ('urgences -> etablissements',
     'urgences.Etablissement',
     'etablissements.nom',
     'N -- 1  (plusieurs patients dans un meme hopital)'),
    ('soins -> etablissements',
     'soins.Etablissement',
     'etablissements.nom',
     'N -- 1  (chaque acte rattache a un hopital)'),
    ('urgences (self)',
     'urgences.IPP',
     'urgences.IPP',
     '1 -- N  (1 patient peut avoir plusieurs visites)'),
])
alt_rows(tbl_rel)

doc.add_paragraph()
h2('Cardinalites')
tbl_card = tbl_grid(4, 3)
header_row(tbl_card, ['Relation', 'Cardinalite', 'Exemple'])
fill_table(tbl_card, [
    ('hopital -- passages',   '1 a N', 'Ibn Sina : 1 hopital -> 287 301 passages'),
    ('passage -- actes',      '1 a N', '1 passage -> 1 a 5 actes medicaux en moyenne'),
    ('patient -- visites',    '1 a N', 'IPP distinct : 505 261 patients, 530 898 passages'),
])
alt_rows(tbl_card)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 6. PIPELINE ETL
# ════════════════════════════════════════════════════════════════════
h1('6. Pipeline ETL')
para(
    'Le pipeline ETL (Extract, Transform, Load) transforme les donnees brutes CSV '
    'en donnees propres et enrichies chargees dans SQLite. '
    'Il suit l\'architecture Medallion en 3 couches : Bronze -> Silver -> Gold.'
)

h2('Architecture Medallion')
tbl_etl = tbl_grid(4, 3)
header_row(tbl_etl, ['Couche', 'Fichiers', 'Description'])
fill_table(tbl_etl, [
    ('Bronze', 'urgences_BRONZE.csv\nsoins_BRONZE.csv\netablissements_BRONZE.csv',
     'Donnees brutes generees par simulation. Pas de nettoyage. Contient doublons et valeurs manquantes.'),
    ('Silver', 'urgences_SILVER.csv\nsoins_SILVER.csv\netablissements_SILVER.csv',
     'Donnees nettoyees : doublons supprimes, dates corrigees, valeurs aberrantes traitees.'),
    ('Gold',   'urgences_GOLD.csv\nsoins_GOLD.csv\nserie_temporelle_daily.csv',
     'Donnees enrichies : nouvelles colonnes calculees (features ML), statistiques agregees.'),
])
alt_rows(tbl_etl)

doc.add_paragraph()
h2('Etape 1 -- Bronze : Lecture')
bullet('Lecture des fichiers CSV avec encodage UTF-8-BOM')
bullet('Normalisation des noms de colonnes (strip, lowercase)')
bullet('Journalisation du nombre de lignes et colonnes')

h2('Etape 2 -- Silver : Nettoyage')
tbl_silver = tbl_grid(8, 2)
header_row(tbl_silver, ['Probleme detecte', 'Traitement applique'])
fill_table(tbl_silver, [
    ('Doublons sur IPP / id_soin',            'drop_duplicates()'),
    ('Dates invalides ou manquantes',          'pd.to_datetime(errors=coerce) + dropna()'),
    ('Duree sejour negative ou > 3 jours',     'Clamp entre 0 et 4320 minutes'),
    ('Age hors [0, 110]',                      'Suppression de la ligne'),
    ('Valeurs manquantes (Antecedents...)',     'Remplacement par valeur par defaut'),
    ('Texte mal formate (majuscules...)',       'str.upper().str.strip()'),
    ('Cout soin negatif',                      'Suppression de la ligne'),
])
alt_rows(tbl_silver)

doc.add_paragraph()
h2('Etape 3 -- Gold : Feature Engineering')
tbl_gold = tbl_grid(11, 3)
header_row(tbl_gold, ['Feature creee', 'Calcul', 'Utilite ML'])
fill_table(tbl_gold, [
    ('Heure_Arrivee',  'date_arrivee.dt.hour',              'Pattern horaire pour XGBoost'),
    ('Jour_Semaine',   'date_arrivee.dt.dayofweek',          'Saisonnalite hebdomadaire'),
    ('Mois / Annee',   'date_arrivee.dt.month / .year',      'Tendance annuelle'),
    ('Saison',         'Regle sur le mois',                  'Saisonnalite climatique'),
    ('Tranche_Horaire','Regle sur heure',                    'Segmentation operationnelle'),
    ('Groupe_Age',     'Regle sur age',                      'Profil patient'),
    ('Est_Pic',        '1 si flux_horaire > Q75 sinon 0',    'Feature critique pour XGBoost'),
    ('Alerte_Charge',  'Regle sur Taux_Hospit + Taux_P1',   'KPI etablissement'),
    ('Serie temporelle','Agregation journaliere des passages','Entree pour Prophet'),
    ('Ratio_Med_Lits', 'nb_medecins / capacite_lits',        'Indicateur RH'),
])
alt_rows(tbl_gold)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 7. BACKEND API
# ════════════════════════════════════════════════════════════════════
h1('7. Backend -- API FastAPI')
para(
    'Le backend expose une API REST documentee automatiquement (Swagger UI sur /docs). '
    'Au demarrage, toutes les donnees sont chargees en memoire RAM en ~40 secondes '
    'via sqlite3 direct pour des reponses < 200ms.'
)

h2('Librairies Python')
tbl_libs = tbl_grid(10, 2)
header_row(tbl_libs, ['Librairie', 'Role'])
fill_table(tbl_libs, [
    ('FastAPI',       'Framework API REST avec documentation automatique Swagger'),
    ('Uvicorn',       'Serveur ASGI haute performance pour FastAPI'),
    ('Pandas',        'Manipulation des DataFrames (530K+ lignes en memoire)'),
    ('NumPy',         'Calculs numeriques et encodages cycliques (sin/cos)'),
    ('SQLAlchemy',    'Connexion SQLite et ORM'),
    ('XGBoost',       'Modele de prediction duree de sejour'),
    ('Scikit-learn',  'Random Forest et preprocessing ML'),
    ('PyJWT',         'Generation et validation des tokens JWT'),
    ('Joblib',        'Chargement des modeles .pkl'),
])
alt_rows(tbl_libs)

doc.add_paragraph()
h2('Endpoints principaux')
tbl_api = tbl_grid(17, 3)
header_row(tbl_api, ['Endpoint', 'Methode', 'Description'])
fill_table(tbl_api, [
    ('/api/auth/login',                  'POST', 'Authentification, retourne token JWT'),
    ('/api/status',                      'GET',  'Etat chargement donnees (ready: true/false)'),
    ('/api/kpis',                        'GET',  'KPIs globaux : total patients, duree moy, taux hospit'),
    ('/api/urgences/orientation',        'GET',  'Repartition par orientation medicale'),
    ('/api/urgences/triage',             'GET',  'Repartition P1/P2/P3/P4'),
    ('/api/urgences/horaire',            'GET',  'Flux horaire moyen par heure'),
    ('/api/urgences/maladies_saisonnieres','GET','Top pathologies par saison'),
    ('/api/etablissements',              'GET',  'Liste des hopitaux avec KPIs calcules'),
    ('/api/predictions',                 'GET',  'Previsions 30 jours par Prophet'),
    ('/api/predictions/planification',   'GET',  'Prevision patients/ressources pour une date'),
    ('/api/simulateur',                  'POST', 'Simulation charge urgences via XGBoost'),
    ('/api/predict/triage',              'POST', 'Evaluation triage patient (P1-P4)'),
    ('/api/modeles/metriques',           'GET',  'Metriques R2, MAE, RMSE des 3 modeles'),
    ('/api/soins/stats',                 'GET',  'Statistiques soins et couts'),
    ('/api/patients',                    'GET',  'Patients actuellement aux urgences'),
    ('/api/patients/{IPP}/statut',       'PATCH','Mettre a jour le statut d\'un patient'),
])
alt_rows(tbl_api)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8. FRONTEND
# ════════════════════════════════════════════════════════════════════
h1('8. Frontend -- Dashboard React')

h2('Pages du dashboard')
tbl_pages = tbl_grid(12, 3)
header_row(tbl_pages, ['Page', 'Route', 'Contenu principal'])
fill_table(tbl_pages, [
    ('KPIs Globaux',          '/',               'Indicateurs cles, alertes, comparaison semaine'),
    ('Analyse Temporelle',    '/temporel',       'Evolution flux, patterns horaires, saisonnalite'),
    ('Etablissements',        '/etablissements', 'Comparaison 8 hopitaux, carte, KPIs detailles'),
    ('Prediction ML',         '/predictions',    'Prophet 30j, simulateur XGBoost, triage, modeles'),
    ('Soins Medicaux',        '/soins',          'Types soins, couts, medicaments, medecins'),
    ('Carte Geographique',    '/carte',          'Carte Maroc avec flux par wilaya'),
    ('Tableau des Donnees',   '/tableau',        'Table complete avec filtres avances et export'),
    ('Alertes & Seuils',      '/alertes',        'Seuils configurables, alertes critiques'),
    ('Gestion Personnel',     '/personnel',      'Planification medecins et infirmiers'),
    ('Gestion Lits',          '/lits',           'Occupation des lits par service'),
    ('Patients Actuels',      '/patients',       'Suivi temps reel des patients presents'),
])
alt_rows(tbl_pages)

doc.add_paragraph()
h2('Bibliotheques frontend')
tbl_front = tbl_grid(7, 2)
header_row(tbl_front, ['Bibliotheque', 'Utilisation'])
fill_table(tbl_front, [
    ('React 18 + TypeScript', 'Framework UI avec typage statique strict'),
    ('Recharts',              'Graphiques : AreaChart, BarChart, RadarChart, PieChart'),
    ('Axios',                 'Appels HTTP vers l\'API avec gestion des tokens JWT'),
    ('React Router v6',       'Navigation SPA entre les pages'),
    ('jsPDF + html2canvas',   'Export des pages en PDF'),
    ('Nginx',                 'Serveur fichiers statiques + proxy /api -> backend:8000'),
])
alt_rows(tbl_front)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 8bis. FONCTIONNEMENT TECHNIQUE DES PAGES DU DASHBOARD
# ════════════════════════════════════════════════════════════════════
h1('8bis. Fonctionnement Technique des Pages du Dashboard')
para(
    'Cette section explique comment chaque page fonctionne reellement : '
    'flux de donnees, formules de calcul, logique de filtrage, appels API, '
    'algorithmes visuels et interactions utilisateur. '
    'Chaque page est un composant React autonome avec son propre cycle de vie.'
)
doc.add_page_break()

# ── Page 1 : KPIs Globaux ────────────────────────────────────────────
h2('Page 1 -- KPIs Globaux  (Route : /)')
para(
    'Page d\'accueil. Au montage du composant React (useEffect), 5 appels API sont lances '
    'en parallele via Promise.all(). Les donnees sont stockees dans des states independants. '
    'Le hook useLive({ intervalMs: 30_000 }) declenche un re-fetch toutes les 30 secondes '
    'via un setInterval interne, et met a jour lastUpdate pour afficher l\'heure du dernier refresh.'
)

doc.add_paragraph()
h3('Flux de donnees au chargement')
code(
    '// Au montage du composant\n'
    'fetchKPIs(\'\', \'\')                  --> GET /api/kpis\n'
    'fetchOrientation()                --> GET /api/urgences/orientation\n'
    'fetchHoraire()                    --> GET /api/urgences/horaire\n'
    'fetchTriage()                     --> GET /api/urgences/triage\n'
    'axios.get(\'/api/stats/comparaison\')--> GET /api/stats/comparaison'
)

doc.add_paragraph()
h3('Calcul du badge TrendBadge (fleche tendance)')
para('Pour chaque KPI de comparaison, un badge colore est calcule :')
code(
    '// delta = valeur_semaine_actuelle - valeur_semaine_precedente\n'
    '// inverse=true  : delta negatif = bon (ex: duree sejour, fugue)\n'
    '// inverse=false : delta positif = bon (ex: nb patients hospitalises)\n'
    'const good = inverse ? val <= 0 : val >= 0;\n'
    'color = val === 0 ? gris : good ? vert (#22C55E) : rouge (#EF4444)\n'
    'arrow = val > 0 ? "up" : "down"'
)

doc.add_paragraph()
h3('Comparaison semaine : structure de donnees retournee par /api/stats/comparaison')
code(
    '{\n'
    '  actuelle:  { patients: N, duree_moy: N, taux_fugue: N, taux_p1: N },\n'
    '  precedente:{ patients: N, duree_moy: N, taux_fugue: N, taux_p1: N },\n'
    '  delta:     { patients: +/-N, duree_moy: +/-N, ... },\n'
    '  periode_actuelle:  "dd/mm/yyyy",\n'
    '  periode_precedente:"dd/mm/yyyy"\n'
    '}'
)

doc.add_paragraph()
h3('Graphiques et leurs donnees')
tbl_p1 = tbl_grid(5, 3)
header_row(tbl_p1, ['Graphique', 'Type Recharts', 'Donnee source'])
fill_table(tbl_p1, [
    ('Orientation patients',   'BarChart vertical',   '/api/urgences/orientation -> [{name, value}]'),
    ('Niveaux de triage',      'PieChart avec Cells', '/api/urgences/triage -> [{name, value}]  -- P1 rouge / P2 orange / P3 bleu / P4 vert'),
    ('Flux horaire 0h-23h',    'BarChart horizontal', '/api/urgences/horaire -> [{heure:0..23, avg_patients:N}]'),
    ('Comparaison semaine',    '4 mini-cards + bars', '/api/stats/comparaison -> delta calcule cote backend'),
])
alt_rows(tbl_p1)
doc.add_page_break()

# ── Page 2 : Analyse Temporelle ─────────────────────────────────────
h2('Page 2 -- Analyse Temporelle  (Route : /temporel)')
para(
    'Le composant maintient un etat annees[] (defaut = toutes les annees 2019-2026). '
    'Chaque modification du filtre annees declenche un useEffect qui relance fetchTemporel(annees.join(",")). '
    'Le parametre est transmis en query string : GET /api/temporel?annees=2022,2023,2024.'
)

doc.add_paragraph()
h3('Algorithme de la moyenne mobile (movingAvg)')
para('Pour lisser la courbe journaliere tres bruitee, un algorithme de moyenne glissante est applique :')
code(
    '// Fenetre = 14 jours (configurable)\n'
    'const movingAvg = (data, window=14) =>\n'
    '  data.map((point, i) => {\n'
    '    const slice = data.slice(Math.max(0, i - window + 1), i + 1);\n'
    '    const avg   = slice.reduce((s, p) => s + p.y, 0) / slice.length;\n'
    '    return { ...point, ma: Math.round(avg * 10) / 10 };\n'
    '  });\n\n'
    '// Puis sous-echantillonnage tous les 7 jours pour reduire les points\n'
    'const tsWithMa = movingAvg(ts, 14).filter((_, i) => i % 7 === 0);'
)
para(
    'Resultat : la serie brute (point par jour) est remplacee par une courbe '
    'lissee (point par semaine), ce qui supprime le bruit quotidien tout en '
    'conservant les tendances mensuelles et saisonnieres.'
)

doc.add_paragraph()
h3('Bascule Serie / Mensuel')
code(
    '// State : view = "serie" | "mensuel"\n'
    '// "serie"   --> AreaChart avec tsWithMa (moving average)\n'
    '// "mensuel" --> BarChart avec agregation mensuelle depuis /api/temporel'
)

doc.add_paragraph()
h3('Statistiques derivees (useMemo)')
para('Les statistiques sont calculees en memoire a partir des donnees chargees, sans nouvel appel API :')
code(
    '// Calculees avec useMemo([ts])\n'
    'totalPatients = ts.reduce((s, p) => s + p.y, 0)\n'
    'maxJour       = Math.max(...ts.map(p => p.y))\n'
    'minJour       = Math.min(...ts.map(p => p.y))\n'
    'tendance      = moyenne(7 derniers jours) - moyenne(7 premiers jours)'
)
doc.add_page_break()

# ── Page 3 : Etablissements ──────────────────────────────────────────
h2('Page 3 -- Etablissements  (Route : /etablissements)')
para(
    'Charge les 8 hopitaux depuis GET /api/etablissements. '
    'Chaque carte est un composant EtabCard. Un clic ouvre EtabModal (modal plein ecran). '
    'Le filtrage se fait en memoire avec useMemo sur le tableau etabs[].'
)

doc.add_paragraph()
h3('Formule du Score de Performance (0 - 100)')
para('Le score est calcule pour chaque hopital par rapport aux 7 autres (normalisation inter-hopitaux) :')
code(
    'const perfScore = (e, all) => {\n'
    '  // norm() : ramene une valeur entre 0 et 100 par rapport au min/max du groupe\n'
    '  const norm = (val, key) => {\n'
    '    const vals = all.map(x => Number(x[key]) || 0);\n'
    '    const mn = Math.min(...vals), mx = Math.max(...vals);\n'
    '    return mx === mn ? 50 : (val - mn) / (mx - mn) * 100;\n'
    '  };\n\n'
    '  const hospit = norm(e.Taux_Hospit_Pct,      "Taux_Hospit_Pct");  // haut = bon\n'
    '  const fugue  = 100 - norm(e.Taux_Fugue_Pct, "Taux_Fugue_Pct");  // bas = bon\n'
    '  const duree  = 100 - norm(e.Duree_Moy_Min,  "Duree_Moy_Min");   // bas = bon\n'
    '  const ratio  = norm(e.Ratio_Medecins_Lits,  "Ratio_Medecins_Lits"); // haut = bon\n\n'
    '  return Math.round((hospit + fugue + duree + ratio) / 4); // moyenne des 4 criteres\n'
    '};\n\n'
    '// Couleur du score :\n'
    '// score >= 70 --> vert  #22C55E\n'
    '// score >= 50 --> orange #F59E0B\n'
    '// score < 50  --> rouge  #EF4444'
)

doc.add_paragraph()
h3('Donnees du RadarChart dans la modal detail')
code(
    '// 5 axes du radar (valeurs en % normalisees 0-100)\n'
    'radarData = [\n'
    '  { metric: "Patients",      val: norm(e.Nb_Patients, ...) },\n'
    '  { metric: "Hospitalisation",val: norm(e.Taux_Hospit_Pct, ...) },\n'
    '  { metric: "Duree sejour",  val: 100 - norm(e.Duree_Moy_Min, ...) },\n'
    '  { metric: "Fugues",        val: 100 - norm(e.Taux_Fugue_Pct, ...) },\n'
    '  { metric: "Qualite",       val: perfScore(e, all) }\n'
    ']'
)

doc.add_paragraph()
h3('Classement Top 3 et medailles')
code(
    '// Tri par score descendant\n'
    'const ranked = [...etabs].sort((a, b) => perfScore(b, etabs) - perfScore(a, etabs));\n'
    '// Les 3 premiers recoivent les medailles Or / Argent / Bronze'
)
doc.add_page_break()

# ── Page 4 : Predictions IA ──────────────────────────────────────────
h2('Page 4 -- Predictions IA  (Route : /predictions)')
para(
    'Page a 6 onglets : Prevision 30j, Saison, Planification, Evaluation Patient, Lits Disponibles, Modeles & Performance. '
    'Chaque onglet est un sous-composant independant monte uniquement quand actif (tab === "prevision" etc.).'
)

doc.add_paragraph()
h3('Onglet Prevision 30 jours -- fonctionnement reel')
para('Appel : GET /api/predictions --> Prophet retourne 30 objets {ds, yhat, yhat_lower, yhat_upper}')
code(
    '// Structure de chaque point Prophet\n'
    'interface Prediction {\n'
    '  ds:          string;  // date "2026-04-15"\n'
    '  yhat:        number;  // prediction centrale\n'
    '  yhat_lower:  number;  // borne basse (IC 80%)\n'
    '  yhat_upper:  number;  // borne haute (IC 80%)\n'
    '}\n\n'
    '// Statistiques calculees en frontend a partir des 30 points\n'
    'avg      = round(sum(yhat) / 30)\n'
    'max      = Math.max(...data.map(p => p.yhat))\n'
    'critical = data.filter(p => p.yhat > avg * 1.2).length  // jours > +20% de la moyenne\n'
    'tendance = moyenne(7 derniers jours) - moyenne(7 premiers jours)'
)
para(
    'Le graphique AreaChart affiche 3 series : "Prevu" (courbe principale), '
    '"Min" (yhat_lower) et "Max" (yhat_upper). Les bornes Min/Max forment '
    'la zone ombree de l\'intervalle de confiance a 80%.'
)

doc.add_paragraph()
h3('Onglet Planification Date -- fonctionnement reel')
para('Appel : GET /api/predictions/planification?date=YYYY-MM-DD')
code(
    '// Le backend retourne pour la date demandee :\n'
    '{\n'
    '  date:            "2026-04-15",\n'
    '  jour_semaine:    "Mardi",\n'
    '  patients_prevus: 183,       // Prophet pour ce jour\n'
    '  medecins_recommandes: 12,   // calcule = patients_prevus / ratio_standard\n'
    '  lits_recommandes: 45,       // calcule = patients_prevus * taux_hospit_moy\n'
    '  historique_meme_jour: [...], // stats des memes jours de semaine historiques\n'
    '  historique_meme_mois: [...], // stats du meme mois historique\n'
    '}'
)

doc.add_paragraph()
h3('Onglet Evaluation Patient -- flux complet')
code(
    '// 1. Utilisateur remplit : age, sexe, motif, antecedents, signes vitaux\n'
    '// 2. POST /api/predict/triage avec le formulaire\n'
    '// 3. Backend : regles expertes + XGBoost -> retourne :\n'
    '{\n'
    '  triage:     "P2",\n'
    '  triage_num: 2,\n'
    '  color:      "#F59E0B",\n'
    '  risque:     "Urgent",\n'
    '  score:      78,\n'
    '  duree_estimee_min: 120\n'
    '}\n'
    '// 4. Frontend affiche le niveau en grand (40px) + barre coloree + duree estimee'
)

doc.add_paragraph()
h3('Onglet Modeles -- Calcul du RadarChart (5 dimensions)')
code(
    '// Metriques recues depuis /api/modeles/metriques :\n'
    '// [ {modele:"Prophet", R2:0.9557, MAE:0.981, RMSE:1.230},\n'
    '//   {modele:"Random Forest", R2:0.9617, MAE:10.514, RMSE:13.266},\n'
    '//   {modele:"XGBoost", R2:0.9670, MAE:9.811, RMSE:12.324} ]\n\n'
    '// Construction des 5 axes du radar :\n'
    'radarData = [\n'
    '  // Axe R2 : R2 * 100\n'
    '  { metric:"R2 Score", Prophet:95.57, "Random Forest":96.17, XGBoost:96.70 },\n'
    '  // Axe Precision : (1 - MAE/maxMAE) * 100\n'
    '  { metric:"Precision", Prophet: (1-0.981/10.514)*100, ... },\n'
    '  // Axe Stabilite : (1 - RMSE/maxRMSE) * 100\n'
    '  { metric:"Stabilite", ... },\n'
    '  // Axe Vitesse : scores fixes (Prophet=62, RF=78, XGBoost=98)\n'
    '  { metric:"Vitesse",   Prophet:62, "Random Forest":78, XGBoost:98 },\n'
    '  // Axe Donnees : scores fixes (Prophet=55, RF=95, XGBoost=95)\n'
    '  { metric:"Donnees",   Prophet:55, "Random Forest":95, XGBoost:95 }\n'
    ']'
)
doc.add_page_break()

# ── Page 5 : Soins Medicaux ──────────────────────────────────────────
h2('Page 5 -- Soins Medicaux  (Route : /soins)')
para(
    'Au chargement, 4 appels API sont lances en parallele via Promise.all(). '
    'Les donnees des 737 512 actes sont aggregees cote backend (SQLite GROUP BY) '
    'avant envoi -- le frontend ne recoit que les Top N necessaires.'
)

doc.add_paragraph()
h3('Flux de chargement')
code(
    'Promise.all([\n'
    '  fetchSoinsTypes(),          // GET /api/soins/types\n'
    '  fetchResultats(),           // GET /api/soins/resultats\n'
    '  fetchMedicaments(),         // GET /api/soins/medicaments\n'
    '  fetchMaladiesSaisonnieres() // GET /api/urgences/maladies_saisonnieres\n'
    '])'
)

doc.add_paragraph()
h3('Structure des donnees recues')
code(
    '// /api/soins/types\n'
    '[{ type_soin: "Consultation", count: 221840, cout_moyen: 350.5 }, ...]\n\n'
    '// /api/soins/resultats\n'
    '[{ resultat: "Gueri", count: 312455, pct: 42.3 }, ...]\n\n'
    '// /api/soins/medicaments\n'
    '[{ medicament: "Paracetamol", count: 98234 }, ...]\n\n'
    '// /api/urgences/maladies_saisonnieres\n'
    '{ Hiver: [{maladie:"Grippe", count:12400}, ...],\n'
    '  Ete:   [{maladie:"Insolation", count:8900}, ...], ... }'
)

doc.add_paragraph()
h3('Affichage des icones saisonniers')
para('Chaque saison est representee par une icone SVG dessinee en React (pas d\'image externe) :')
code(
    '// Hiver   --> IconSnowflake  (flocon de neige, 6 branches)\n'
    '// Printemps --> IconFlower  (fleur 6 petales)\n'
    '// Ete     --> IconSun      (soleil avec rayons)\n'
    '// Automne --> IconLeaf     (feuille)\n'
    '// Couleurs: Hiver=#3B82F6, Printemps=#22C55E, Ete=#F59E0B, Automne=#F97316'
)
doc.add_page_break()

# ── Page 6 : Carte Geographique ──────────────────────────────────────
h2('Page 6 -- Carte Geographique  (Route : /carte)')
para(
    'Utilise react-leaflet (MapContainer + TileLayer + CircleMarker). '
    'Les coordonnees GPS (lat/lng) de chaque hopital sont stockees en base '
    'et retournees par GET /api/etablissements.'
)

doc.add_paragraph()
h3('Calcul du rayon des marqueurs')
code(
    '// Rayon proportionnel au nombre de patients de l\'hopital\n'
    '// par rapport au maximum du reseau\n'
    'const maxPatients = Math.max(...etabs.map(e => e.Nb_Patients));\n'
    'const radius = 8 + (e.Nb_Patients / maxPatients) * 22;\n'
    '// Hopital le plus charge: radius=30  |  Hopital le moins charge: radius=8'
)

doc.add_paragraph()
h3('Couleur des marqueurs selon l\'alerte')
code(
    'ALERTE_COLOR = {\n'
    '  Normal:   "#22C55E",  // vert\n'
    '  Eleve:    "#F59E0B",  // orange\n'
    '  Critique: "#EF4444",  // rouge\n'
    '}\n'
    '// La valeur Alerte_Charge est calculee en ETL :\n'
    '// Critique si Taux_P1 > 10% OU Taux_Fugue > 8%\n'
    '// Eleve    si Taux_P1 > 5%  OU Taux_Fugue > 4%\n'
    '// Normal   sinon'
)

doc.add_paragraph()
h3('Interaction : clic sur marqueur / carte')
code(
    '// useState<EtabGeo | null> selected\n'
    '// CircleMarker onClick --> setSelected(etab)\n'
    '// Le panneau lateral reagit au selected et affiche les StatBox\n'
    '// StatBox : label + valeur numerique dans une mini-carte'
)
doc.add_page_break()

# ── Page 7 : Tableau des Donnees ─────────────────────────────────────
h2('Page 7 -- Tableau des Donnees  (Route : /tableau)')
para(
    'Charge toutes les donnees en memoire RAM via GET /api/tableau (sans pagination serveur). '
    'Le filtrage, le tri et la pagination sont entierement geres en frontend avec useMemo, '
    'ce qui rend les filtres instantanes sans nouvel appel API.'
)

doc.add_paragraph()
h3('Pipeline de filtrage (useMemo)')
code(
    '// Le tableau filtered est recalcule a chaque changement de filtre\n'
    'const filtered = useMemo(() => {\n'
    '  let d = [...rows];\n'
    '  // Recherche texte (nom, etablissement, triage, orientation, id_passage)\n'
    '  if (search)\n'
    '    d = d.filter(r =>\n'
    '      r.Nom_Complet.toLowerCase().includes(search.toLowerCase()) ||\n'
    '      r.Etablissement.toLowerCase().includes(search.toLowerCase()) ||\n'
    '      String(r.id_passage).includes(search)\n'
    '    );\n'
    '  if (filterTriage) d = d.filter(r => r.Niveau_Triage === filterTriage);\n'
    '  if (filterOrient) d = d.filter(r => r.Orientation    === filterOrient);\n'
    '  if (filterAnnee)  d = d.filter(r => String(r.Annee)  === filterAnnee);\n'
    '  if (filterSexe)   d = d.filter(r => r.Sexe           === filterSexe);\n'
    '  if (filterCIN)    d = d.filter(r => r.CIN            === filterCIN);\n'
    '  // Tri par colonne (sortKey / sortDir)\n'
    '  d.sort((a, b) => sortDir === "asc" ? a[sortKey] > b[sortKey] ? 1 : -1\n'
    '                                      : a[sortKey] < b[sortKey] ? 1 : -1);\n'
    '  return d;\n'
    '}, [rows, search, filterTriage, filterOrient, filterAnnee, filterSexe, filterCIN,\n'
    '    sortKey, sortDir]);'
)

doc.add_paragraph()
h3('Pagination front-end')
code(
    'const PAGE_SIZE  = 50; // lignes par page\n'
    'const totalPages = Math.ceil(filtered.length / PAGE_SIZE);\n'
    'const pageData   = filtered.slice((page-1)*PAGE_SIZE, page*PAGE_SIZE);\n\n'
    '// Boutons de pagination : affiche max 5 numeros autour de la page courante\n'
    'const delta = 2;\n'
    'const start = Math.max(1, Math.min(page - delta, totalPages - delta*2));\n'
    'const end   = Math.min(totalPages, Math.max(page + delta, delta*2 + 1));'
)

doc.add_paragraph()
h3('Statistiques calculees sur les donnees filtrees')
code(
    '// Recalculees a chaque changement de filtre (useMemo)\n'
    'const stats = {\n'
    '  avgAge:  Math.round(filtered.reduce((s,r) => s + r.Age, 0) / filtered.length),\n'
    '  avgDur:  Math.round(filtered.reduce((s,r) => s + r.Duree_Sejour_min, 0) / filtered.length),\n'
    '  p1count: filtered.filter(r => r.Niveau_Triage.startsWith("P1")).length,\n'
    '  hospit:  filtered.filter(r => r.Orientation === "Hospitalise").length,\n'
    '  total:   filtered.length\n'
    '};'
)

doc.add_paragraph()
h3('WordCloud des antecedents')
code(
    '// Compte les occurrences de chaque antecedent dans les donnees filtrees\n'
    'const antecedentsWords = useMemo(() => {\n'
    '  const counts = {};\n'
    '  filtered.forEach(r => {\n'
    '    const word = r.Antecedents || "Aucun";\n'
    '    counts[word] = (counts[word] || 0) + 1;\n'
    '  });\n'
    '  return Object.entries(counts).map(([text, value]) => ({ text, value }));\n'
    '}, [filtered]);\n'
    '// Passe a <WordCloud words={antecedentsWords} /> qui calcule la taille\n'
    '// de chaque mot proportionnellement a sa frequence'
)
doc.add_page_break()

# ── Page 8 : Alertes & Seuils ────────────────────────────────────────
h2('Page 8 -- Alertes et Seuils  (Route : /alertes)')
para(
    'Deux etapes au chargement : lecture de la config actuelle, puis verification immediate. '
    'A chaque sauvegarde, la config est envoyee au backend qui la persiste en JSON, '
    'puis une nouvelle verification est faite et le resultat re-affiche instantanement.'
)

doc.add_paragraph()
h3('Structure de la configuration (valeurs par defaut OMS)')
code(
    'interface AlertConfig {\n'
    '  duree_moy_seuil:   number;  // defaut 240 min (objectif OMS)\n'
    '  taux_fugue_seuil:  number;  // defaut 3%  (recommandation nationale)\n'
    '  taux_p1_seuil:     number;  // defaut 5%  (seuil critique P1)\n'
    '  taux_hospit_seuil: number;  // defaut 35% (seuil saturation)\n'
    '}'
)

doc.add_paragraph()
h3('Flux de sauvegarde et re-verification')
code(
    'const handleSave = async () => {\n'
    '  // 1. Envoie la nouvelle config\n'
    '  await axios.post(\'/api/alertes/config\', config);\n'
    '  // 2. Re-verifie immediatement les alertes avec la nouvelle config\n'
    '  const r = await axios.get(\'/api/alertes/check\');\n'
    '  setResult(r.data);\n'
    '  // 3. Toast de confirmation vert (via ToastContext)\n'
    '  showToast({ type: "success", title: "Seuils mis a jour", ... });\n'
    '};'
)

doc.add_paragraph()
h3('Structure d\'une alerte retournee par /api/alertes/check')
code(
    'interface Alerte {\n'
    '  type:   string;          // ex: "Duree sejour"\n'
    '  valeur: string;          // ex: "267 min"\n'
    '  seuil:  string;          // ex: "240 min"\n'
    '  niveau: "critique" | "warning";\n'
    '}\n\n'
    '// Logique d\'affichage :\n'
    '// niveau="critique" --> fond rouge #fee2e2, bordure gauche #ef4444\n'
    '// niveau="warning"  --> fond jaune #fef3c7, bordure gauche #f59e0b'
)
doc.add_page_break()

# ── Page 9 : Admission Patient ───────────────────────────────────────
h2('Page 9 -- Admission Patient  (Route : /admission)')
para(
    'Le formulaire utilise un seul state "form" de type objet, '
    'pre-initialise avec des valeurs par defaut. '
    'La date d\'arrivee est pre-remplie avec l\'heure exacte du navigateur. '
    'L\'indicateur de triage se met a jour visuellement a chaque changement du select.'
)

doc.add_paragraph()
h3('Etat initial du formulaire')
code(
    'const initialForm = {\n'
    '  nom_complet: \'\', cin: \'\', age: \'\', sexe: \'M\',\n'
    '  groupe_sanguin: \'O+\', antecedents: \'Aucun\',\n'
    '  etablissement: \'Hopital Ibn Sina\',\n'
    '  niveau_triage: \'P3 - Semi-urgent\',\n'
    '  motif_consultation: \'Malaise\',\n'
    '  orientation: \'Domicile\',\n'
    '  duree_sejour_min: \'60\',\n'
    '  date_arrivee: new Date().toISOString().slice(0, 16), // "2026-04-14T10:35"\n'
    '  mutuelle: \'Payant\', prix_sejour: \'0\', prix_soins: \'0\',\n'
    '};'
)

doc.add_paragraph()
h3('Flux de soumission')
code(
    '// 1. Validation : nom_complet et cin obligatoires\n'
    '// 2. POST /api/admission avec le formulaire complet\n'
    '// 3. Backend : insere dans urgences, genere un IPP unique\n'
    '// 4. Backend retourne : { IPP: "505262" }\n'
    '// 5. Frontend :\n'
    '//    - setResult({ IPP: res.data.IPP })\n'
    '//    - Affiche la carte de confirmation avec l\'IPP\n'
    '//    - Toast vert : "Patient admis -- IPP: 505262"\n'
    '//    - setForm(initialForm) --> reinitialise tout le formulaire'
)

doc.add_paragraph()
h3('Indicateur visuel du triage')
code(
    'NIVEAU_COLORS = {\n'
    '  "P1 - Critique":     "#ef4444",  // rouge\n'
    '  "P2 - Urgent":       "#f97316",  // orange\n'
    '  "P3 - Semi-urgent":  "#f59e0b",  // jaune\n'
    '  "P4 - Non urgent":   "#22c55e",  // vert\n'
    '}\n'
    '// Une barre coloree sous le select change de couleur en temps reel\n'
    '// selon la valeur selectionnee dans form.niveau_triage'
)
doc.add_page_break()

# ── Page 10 : Patients Actuels ───────────────────────────────────────
h2('Page 10 -- Patients Actuels  (Route : /patients)')
para(
    'Utilise le hook useAutoRefresh(load, 30000) qui appelle load() '
    'toutes les 30 secondes via setInterval pour mettre a jour la liste des patients. '
    'La mise a jour du statut d\'un patient est optimiste : le state React est '
    'mis a jour immediatement, sans attendre la confirmation du serveur.'
)

doc.add_paragraph()
h3('Calcul du taux de charge en temps reel')
code(
    '// Retourne par /api/patients (cote backend)\n'
    'interface LiveKPIs {\n'
    '  patients_aujourd_hui: number,  // admis dans les dernieres 24h\n'
    '  patients_actifs:      number,  // statut != "Sorti"\n'
    '  lits_occupes:         number,\n'
    '  total_lits:           number,\n'
    '  taux_charge:          number,  // lits_occupes / total_lits * 100\n'
    '  taux_p1_aujourd_hui:  number,  // % P1 parmi les patients du jour\n'
    '}\n\n'
    '// Couleur de la barre de charge :\n'
    '// taux_charge > 80% --> rouge  #EF4444\n'
    '// taux_charge > 60% --> orange #F59E0B\n'
    '// sinon            --> vert   #22C55E'
)

doc.add_paragraph()
h3('Mise a jour du statut patient')
code(
    'const updateStatut = async (IPP, statut, lit_numero?) => {\n'
    '  // Mise a jour optimiste du state React (affichage immediat)\n'
    '  setPatients(prev => prev.map(p =>\n'
    '    p.IPP === IPP ? { ...p, statut, lit_numero: lit_numero ?? p.lit_numero } : p\n'
    '  ));\n'
    '  // Appel API asynchrone\n'
    '  await axios.patch(`/api/patients/${IPP}/statut`, {\n'
    '    statut,\n'
    '    lit_numero,\n'
    '    updated_by: currentUser.username\n'
    '  });\n'
    '  // Le backend enregistre updated_at et updated_by dans patient_statuts\n'
    '};'
)

doc.add_paragraph()
h3('Compteurs par statut (affichage en bas de page)')
code(
    'const statCounts = STATUTS.reduce(\n'
    '  (acc, s) => ({ ...acc, [s]: patients.filter(p => p.statut === s).length }),\n'
    '  {}\n'
    ');\n'
    '// STATUTS = ["En triage", "En attente", "En traitement", "Sorti"]'
)
doc.add_page_break()

# ── Page 11 : Gestion Personnel ─────────────────────────────────────
h2('Page 11 -- Gestion du Personnel  (Route : /personnel)')
para(
    'Le composant charge GET /api/personnel et groupe les agents par etablissement '
    'en memoire (Object.groupBy ou reduce). Chaque groupe est un accordeon '
    'independant avec un state open/closed par etablissement.'
)

doc.add_paragraph()
h3('Groupement du personnel par etablissement')
code(
    '// Apres reception de la liste flat depuis /api/personnel :\n'
    'const grouped = personnel.reduce((acc, p) => {\n'
    '  if (!acc[p.etablissement]) acc[p.etablissement] = [];\n'
    '  acc[p.etablissement].push(p);\n'
    '  return acc;\n'
    '}, {} as Record<string, Personnel[]>);\n\n'
    '// Chaque cle est un nom d\'hopital :\n'
    '// { "Hopital Ibn Sina": [...], "CHU Rabat": [...], ... }'
)

doc.add_paragraph()
h3('Couleurs des statuts')
code(
    'STATUT_COLORS = {\n'
    '  "En service": { bg:"#D1FAE5", color:"#065F46", border:"#10B981" }, // vert\n'
    '  "En garde":   { bg:"#FEF3C7", color:"#D97706", border:"#F59E0B" }, // jaune\n'
    '  "En conge":   { bg:"#DBEAFE", color:"#1D4ED8", border:"#3B82F6" }, // bleu\n'
    '  "Repos":      { bg:"#F1F5F9", color:"#64748B", border:"#94A3B8" }, // gris\n'
    '}'
)

doc.add_paragraph()
h3('Mise a jour du statut d\'un agent')
code(
    '// Dropdown dans la ligne de l\'agent\n'
    '// onChange --> axios.patch(`/api/personnel/${agent.id}/statut`, { statut })\n'
    '// --> mise a jour en base SQLite dans la table personnel\n'
    '// --> updated_at et updated_by enregistres automatiquement'
)
doc.add_page_break()

# ── Page 12 : Gestion Lits ──────────────────────────────────────────
h2('Page 12 -- Gestion des Lits  (Route : /lits)')
para(
    'Meme architecture que la page Personnel : chargement flat, groupement par etablissement, '
    'accordeons independants. Le composant StatBar calcule et affiche la repartition '
    'des statuts en une mini-barre coloree par etablissement.'
)

doc.add_paragraph()
h3('Composant StatBar -- calcul de la barre de repartition')
code(
    '// Pour un etablissement donne, calcule la proportion de chaque statut\n'
    'const StatBar = ({ lits }) => {\n'
    '  const total = lits.length;\n'
    '  // Affiche 4 segments colores proportionnels\n'
    '  return STATUTS_LIT.map(statut => {\n'
    '    const count = lits.filter(l => l.statut === statut).length;\n'
    '    const width = `${(count / total) * 100}%`;\n'
    '    return <div style={{ width, background: STATUT_COLORS[statut] }} />;\n'
    '  });\n'
    '};\n\n'
    'STATUTS_LIT = ["Disponible","Occupe","En maintenance","Reserve"]\n'
    '// Couleurs :\n'
    '// Disponible     --> vert  #10B981\n'
    '// Occupe         --> rouge #EF4444\n'
    '// En maintenance --> orange #F59E0B\n'
    '// Reserve        --> bleu  #3B82F6'
)

doc.add_paragraph()
h3('Mise a jour du statut d\'un lit')
code(
    '// Clic sur le badge statut d\'un lit --> dropdown\n'
    '// onChange --> axios.patch(`/api/lits/${lit.id}/statut`, {\n'
    '//   statut: "Disponible",    // ou Occupe / En maintenance / Reserve\n'
    '//   nom_patient: "...",      // si statut = Occupe\n'
    '//   date_admission: "..."    // si statut = Occupe\n'
    '// })\n'
    '// Le lit est mis a jour dans la table lits de chu.db'
)
doc.add_page_break()

# ── Page 13 : Vue d'ensemble RH ─────────────────────────────────────
h2('Page 13 -- Vue d\'ensemble RH  (Route : /ressources)')
para(
    'Cette page est une vue agregee qui combine 3 sources de donnees : '
    'la liste du personnel, l\'occupation des lits, et les informations des etablissements. '
    'Le backend construit une jointure entre ces tables et retourne un objet '
    'EtabRow par hopital contenant toutes les metriques.'
)

doc.add_paragraph()
h3('Structure EtabRow retournee par /api/ressources/vue_ensemble')
code(
    'interface EtabRow {\n'
    '  etablissement:    string;\n'
    '  directeur:        Directeur | null;  // agent avec role="Directeur"\n'
    '  medecins:         number;            // count role="Medecin"\n'
    '  infirmiers:       number;            // count role="Infirmier"\n'
    '  en_service:       number;            // count statut="En service"\n'
    '  lits_total:       number;\n'
    '  lits_dispo:       number;\n'
    '  lits_occupes:     number;\n'
    '  lits_maintenance: number;\n'
    '  taux_occupation:  number;  // lits_occupes / lits_total * 100\n'
    '}'
)

doc.add_paragraph()
h3('Requete SQL de construction (cote backend)')
code(
    '-- Agregation personnel par etablissement\n'
    'SELECT etablissement,\n'
    '       COUNT(*) FILTER (WHERE role="Medecin")         AS medecins,\n'
    '       COUNT(*) FILTER (WHERE role="Infirmier")       AS infirmiers,\n'
    '       COUNT(*) FILTER (WHERE statut="En service")    AS en_service\n'
    'FROM personnel GROUP BY etablissement;\n\n'
    '-- Agregation lits par etablissement\n'
    'SELECT etablissement,\n'
    '       COUNT(*) AS total,\n'
    '       COUNT(*) FILTER (WHERE statut="Disponible")    AS dispo,\n'
    '       COUNT(*) FILTER (WHERE statut="Occupe")        AS occupes\n'
    'FROM lits GROUP BY etablissement;\n\n'
    '-- Jointure Python (pandas merge) sur le nom de l\'etablissement'
)

doc.add_paragraph()
h3('Alerte RH automatique')
code(
    '// Affichee pour chaque hopital en rouge si :\n'
    'if (row.taux_occupation > 90) --> "Saturation lits"\n'
    'if (row.en_service / (row.medecins + row.infirmiers) < 0.6)\n'
    '  --> "Personnel insuffisant en service"'
)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 9. MODÈLES ML
# ════════════════════════════════════════════════════════════════════
h1('9. Modeles Machine Learning')
para(
    'Trois modeles ML ont ete entraines sur les donnees historiques (2019-2024) '
    'et evalues sur les donnees 2025-2026 (20% test = 106 180 lignes).'
)

h2('9.1  Prophet -- Prevision de l\'affluence')
tbl_prophet = tbl_grid(6, 2)
header_row(tbl_prophet, ['Parametre', 'Valeur / Description'])
fill_table(tbl_prophet, [
    ('Type',        'Modele de series temporelles additif (Meta/Facebook)'),
    ('Entree',      'Serie journaliere : date + nb patients (2019-2026, ~2 500 points)'),
    ('Sortie',      'Previsions J+1 a J+30 avec intervalle de confiance 80%'),
    ('Formule',     'y(t) = g(t) + s(t) + h(t) + erreur'),
    ('Composantes', 'g(t)=tendance  s(t)=saisonnalite (Fourier)  h(t)=jours feries'),
])
alt_rows(tbl_prophet)

doc.add_paragraph()
h2('9.2  Random Forest -- Duree de sejour')
tbl_rf = tbl_grid(6, 2)
header_row(tbl_rf, ['Parametre', 'Valeur / Description'])
fill_table(tbl_rf, [
    ('Type',        'Ensemble de 500 arbres de decision (bagging)'),
    ('Entree',      '27 features : age, triage, heure, saison, lits, medecins...'),
    ('Sortie',      'Duree estimee du sejour en minutes'),
    ('Principe',    'Prediction = moyenne des 500 arbres (reduit le surapprentissage)'),
    ('Feature cle', 'Triage (P1-P4) et age sont les variables les plus predictives'),
])
alt_rows(tbl_rf)

doc.add_paragraph()
h2('9.3  XGBoost -- Duree de sejour (optimal)')
tbl_xgb = tbl_grid(6, 2)
header_row(tbl_xgb, ['Parametre', 'Valeur / Description'])
fill_table(tbl_xgb, [
    ('Type',         'Gradient Boosting avec regularisation L1/L2'),
    ('Entree',       '27 features + encodages cycliques sin/cos heure et jour'),
    ('Sortie',       'Duree estimee du sejour en minutes'),
    ('Principe',     '~100 arbres correcteurs en sequence, chacun corrige les residus'),
    ('Performance',  'Prediction < 1ms par patient -- utilise en temps reel'),
])
alt_rows(tbl_xgb)

doc.add_paragraph()
h2('Resultats comparatifs')
tbl_results = tbl_grid(4, 6)
header_row(tbl_results, ['Modele', 'Type', 'R2 Score', 'MAE', 'RMSE', 'Utilisation'])
fill_table(tbl_results, [
    ('Prophet',         'Serie temporelle',  '95.57%', '0.981 patient/j', '1.230', 'Prevision affluence 30j'),
    ('Random Forest',   'Ensembliste',        '96.17%', '10.514 min',      '13.266 min', 'Duree sejour'),
    ('XGBoost (best)',  'Gradient Boosting',  '96.70%', '9.811 min',       '12.324 min', 'Simulateur + triage'),
])
alt_rows(tbl_results)
for cell in tbl_results.rows[3].cells:
    shade(cell, 'FFF2CC')
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 10. FONCTIONNALITÉS
# ════════════════════════════════════════════════════════════════════
h1('10. Fonctionnalites du Dashboard')
tbl_feat2 = tbl_grid(13, 2)
header_row(tbl_feat2, ['Fonctionnalite', 'Description technique'])
fill_table(tbl_feat2, [
    ('Mode sombre / clair',      'ThemeContext React, preference persistee en localStorage'),
    ('Export PDF',               'html2canvas capture + jsPDF generation, telechargement direct'),
    ('Envoi email',              'Rapport PDF envoye par SMTP depuis l\'interface (EmailModal)'),
    ('Filtres globaux',          'Par annee, etablissement, orientation -- appliques a tous les graphiques'),
    ('Live refresh KPIs',        'Rafraichissement automatique toutes les 30 secondes (hook useLive)'),
    ('Simulateur ML',            'Formulaire -> POST /api/simulateur -> XGBoost -> prediction charge'),
    ('Evaluation triage',        'Profil patient -> regles expertes + XGBoost -> niveau P1-P4 + duree'),
    ('Planification date',       'Date choisie -> historique jour/mois -> patients + medecins + lits'),
    ('Ecran de chargement',      'BackendLoader poll /api/status toutes les 3s jusqu\'a ready:true'),
    ('Gestion des roles',        'JWT payload role -> canAccess() -> redirect ou affichage conditionnel'),
    ('Graphique radar ML',       'RadarChart Recharts -- 5 dimensions : R2, Precision, Stabilite, Vitesse'),
    ('Algo depliable (ML)',       'Accordeon par modele avec 5 etapes detaillees du fonctionnement'),
])
alt_rows(tbl_feat2)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 11. SÉCURITÉ
# ════════════════════════════════════════════════════════════════════
h1('11. Securite et Authentification')

h2('JWT (JSON Web Token)')
bullet('Algorithme de signature : HS256')
bullet('Duree de validite configurable : JWT_EXPIRE_MINUTES (defaut 480 min = 8h)')
bullet('Toutes les routes protegees par HTTPBearer dependency FastAPI')
bullet('Token stocke en localStorage cote client')

doc.add_paragraph()
h2('Systeme de roles')
tbl_roles = tbl_grid(5, 3)
header_row(tbl_roles, ['Role', 'Pages accessibles', 'Droits speciaux'])
fill_table(tbl_roles, [
    ('admin',     'Toutes les pages',                              'Gestion utilisateurs, toutes les actions'),
    ('medecin',   'KPIs, Analyse, Predictions, Patients, Soins',  'Evaluation triage, admission patient'),
    ('infirmier', 'KPIs, Patients actuels, Gestion lits',         'Mise a jour statut patient'),
    ('direction', 'KPIs, Etablissements, Soins, Personnel',       'Export PDF, statistiques globales'),
])
alt_rows(tbl_roles)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 12. DÉPLOIEMENT
# ════════════════════════════════════════════════════════════════════
h1('12. Deploiement Docker')

h2('Commandes')
code(
    '# Premier demarrage complet\n'
    'docker-compose up -d\n\n'
    '# Rebuild frontend uniquement (ne redemmarre pas le backend)\n'
    'docker-compose build frontend && docker-compose up -d frontend\n\n'
    '# Rebuild backend uniquement\n'
    'docker-compose build backend && docker-compose up -d backend\n\n'
    '# Voir les logs en direct\n'
    'docker logs chu_backend --follow\n\n'
    '# Verifier l\'etat du chargement des donnees\n'
    'curl http://localhost:8000/api/status'
)

doc.add_paragraph()
h2('Variables d\'environnement (.env)')
tbl_env = tbl_grid(8, 3)
header_row(tbl_env, ['Variable', 'Valeur par defaut', 'Description'])
fill_table(tbl_env, [
    ('DATABASE_URL',        'sqlite:////app/data/chu.db', 'Chemin de la base de donnees'),
    ('JWT_SECRET',          '(obligatoire)',               'Cle secrete pour signer les tokens JWT'),
    ('JWT_EXPIRE_MINUTES',  '480',                         'Duree de validite du token (8 heures)'),
    ('SMTP_HOST',           '---',                         'Serveur email pour envoi des rapports'),
    ('SMTP_PORT',           '587',                         'Port SMTP (TLS)'),
    ('SMTP_USER',           '---',                         'Adresse email expediteur'),
    ('CORS_ORIGINS',        '*',                           'Origines autorisees pour les requetes CORS'),
])
alt_rows(tbl_env)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# 13. PERFORMANCES
# ════════════════════════════════════════════════════════════════════
h1('13. Performances et Resultats')

tbl_perf = tbl_grid(8, 2)
header_row(tbl_perf, ['Metrique', 'Valeur mesuree'])
fill_table(tbl_perf, [
    ('Temps demarrage backend (chargement donnees)', '~40 secondes'),
    ('Temps de reponse API (requetes analytics)',     '< 200 ms'),
    ('Temps de prediction XGBoost (par patient)',     '< 1 ms'),
    ('Taille base de donnees chu.db',                '~350 MB'),
    ('Memoire RAM utilisee (backend + donnees)',     '~1.2 GB'),
    ('Score R2 meilleur modele (XGBoost)',            '96.70%'),
    ('Score R2 Prophet (prevision 30 jours)',         '95.57%'),
])
alt_rows(tbl_perf)

doc.add_paragraph()
h1('Conclusion')
para(
    'Ce projet de fin d\'etudes demontre la faisabilite d\'un tableau de bord '
    'decisionnel complet pour les urgences hospitalieres, integrant l\'analyse '
    'de donnees historiques, la visualisation interactive et la prediction par '
    'Machine Learning dans une architecture moderne et containerisee. '
    'Les modeles atteignent un R2 superieur a 95% sur l\'ensemble de test, '
    'validant leur pertinence pour une utilisation en production. '
    'L\'architecture Docker garantit la portabilite et la facilite de deploiement '
    'sur n\'importe quel environnement serveur.'
)

# ── Sauvegarde ───────────────────────────────────────────────────────
out = r'C:\Users\Azddine\Desktop\PFE\Documentation_Finale_PFE_v2.docx'
doc.save(out)
print(f'Document genere : {out}')
