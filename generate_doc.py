"""
Génère le document récapitulatif Word du PFE CHU Ibn Sina
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Couleurs ─────────────────────────────────────────────────
BLUE_DARK   = RGBColor(0x0F, 0x17, 0x29)   # #0f1729
BLUE_MAIN   = RGBColor(0x1a, 0x31, 0x70)   # #1a3170
BLUE_LIGHT  = RGBColor(0x3b, 0x82, 0xf6)   # #3b82f6
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_LIGHT  = RGBColor(0xF3, 0xF4, 0xF6)
GRAY_MID    = RGBColor(0x6B, 0x72, 0x80)
GREEN       = RGBColor(0x16, 0xa3, 0x4a)
RED         = RGBColor(0xDC, 0x26, 0x26)
ORANGE      = RGBColor(0xEA, 0x58, 0x0C)

doc = Document()

# ── Marges ───────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helpers ───────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'),   'single')
        border.set(qn('w:sz'),    '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_run(para, txt, bold=False, italic=False, size=11,
            color=None, font="Calibri"):
    run = para.add_run(txt)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def heading(level, txt, color=BLUE_MAIN, size=None):
    sizes = {1:22, 2:16, 3:13, 4:12}
    sz    = size or sizes.get(level, 12)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level==1 else 8)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(txt)
    run.bold = True
    run.font.name  = "Calibri"
    run.font.size  = Pt(sz)
    run.font.color.rgb = color
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        # underline separator
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '6')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), '1a3170')
        pBdr.append(bot)
        pPr.append(pBdr)
    return p

def body(txt, size=11, color=None, bold=False, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    add_run(p, txt, bold=bold, size=size, color=color)
    return p

def bullet(txt, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=10.5)
        add_run(p, txt, size=10.5)
    else:
        add_run(p, txt, size=10.5)
    return p

def make_table(headers, rows, col_widths=None, header_bg="1a3170"):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row in enumerate(rows):
        bg = "F9FAFB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def code_block(txt):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F3F4F6')
    pPr.append(shd)
    run = p.add_run(txt)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    return p

def page_break():
    doc.add_page_break()

# ════════════════════════════════════════════════════════════
#  PAGE DE GARDE
# ════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("CHU IBN SINA — RABAT")
run.bold = True
run.font.name  = "Calibri"
run.font.size  = Pt(14)
run.font.color.rgb = BLUE_MAIN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(40)
run2 = p2.add_run("TABLEAU DE BORD DÉCISIONNEL\nPOUR LA GESTION DES URGENCES HOSPITALIÈRES")
run2.bold = True
run2.font.name  = "Calibri"
run2.font.size  = Pt(22)
run2.font.color.rgb = BLUE_DARK

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(8)
run3 = p3.add_run("Document Récapitulatif du Projet")
run3.font.name  = "Calibri"
run3.font.size  = Pt(14)
run3.font.color.rgb = BLUE_LIGHT
run3.italic = True

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_rows = [
    ("Auteur",       "Azddine"),
    ("Filière",      "BI & Data Science"),
    ("Année",        "2024 / 2025"),
    ("Encadrement",  "CHU Ibn Sina, Rabat"),
    ("Date",         "Mai 2026"),
]
for i, (k, v) in enumerate(info_rows):
    c0, c1 = info_table.rows[i].cells[0], info_table.rows[i].cells[1]
    set_cell_bg(c0, "E8EEF8")
    set_cell_bg(c1, "FFFFFF")
    p0 = c0.paragraphs[0]
    p1 = c1.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r0 = p0.add_run(k + "  ")
    r0.bold = True; r0.font.size = Pt(11); r0.font.name = "Calibri"
    r0.font.color.rgb = BLUE_MAIN
    r1 = p1.add_run("  " + v)
    r1.font.size = Pt(11); r1.font.name = "Calibri"
    r1.font.color.rgb = RGBColor(0x11, 0x18, 0x27)
    c0.width = Cm(5); c1.width = Cm(8)

page_break()

# ════════════════════════════════════════════════════════════
#  SOMMAIRE
# ════════════════════════════════════════════════════════════
heading(1, "Sommaire")
toc_items = [
    ("1.", "Contexte et Objectif"),
    ("2.", "Stack Technique"),
    ("3.", "Architecture du Projet"),
    ("4.", "Base de Données"),
    ("5.", "Pages du Tableau de Bord"),
    ("6.", "Modèles Machine Learning"),
    ("7.", "Fonctionnalités Clés"),
    ("8.", "API — Principaux Endpoints"),
    ("9.", "Déploiement"),
    ("10.", "Problèmes Résolus"),
    ("11.", "Conclusion"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, f"  {num}  ", bold=True, size=11, color=BLUE_MAIN)
    add_run(p, title, size=11)

page_break()

# ════════════════════════════════════════════════════════════
#  1. CONTEXTE ET OBJECTIF
# ════════════════════════════════════════════════════════════
heading(1, "1.  Contexte et Objectif")

body(
    "Le projet de fin d'études (PFE) consiste à concevoir et développer un tableau de bord "
    "décisionnel complet pour optimiser la gestion des urgences du CHU Ibn Sina de Rabat. "
    "Il s'inscrit dans le cadre de la filière BI & Data Science et couvre l'ensemble de la "
    "chaîne de valeur : de la collecte et la préparation des données jusqu'à la mise en "
    "production d'une application web full-stack."
)

heading(2, "Objectifs principaux")
bullet("Anticiper les flux de patients grâce aux modèles de prédiction (ML / séries temporelles)")
bullet("Réduire les temps d'attente aux urgences par une meilleure allocation des ressources")
bullet("Fournir un suivi en temps réel des patients admis (statut, lit attribué)")
bullet("Centraliser les données de 8 établissements du réseau CHU")
bullet("Offrir des alertes automatiques sur les indicateurs critiques")
bullet("Permettre la planification RH et lits pour une date future donnée")

page_break()

# ════════════════════════════════════════════════════════════
#  2. STACK TECHNIQUE
# ════════════════════════════════════════════════════════════
heading(1, "2.  Stack Technique")

heading(2, "Backend")
make_table(
    ["Technologie", "Version", "Rôle"],
    [
        ("Python",       "3.11",   "Langage principal"),
        ("FastAPI",      "latest", "API REST — port 8000"),
        ("SQLAlchemy",   "2.x",    "ORM multi-base de données"),
        ("PyMySQL",      "latest", "Driver MySQL"),
        ("pandas",       "3.0+",   "Traitement et analyse des données"),
        ("numpy",        "latest", "Calculs numériques"),
        ("XGBoost",      "latest", "Modèle de prédiction ML"),
        ("scikit-learn", "latest", "Pipeline ML, métriques"),
        ("joblib",       "latest", "Sérialisation des modèles .pkl"),
        ("PyJWT",        "latest", "Authentification par tokens JWT"),
        ("openpyxl",     "latest", "Export Excel (.xlsx)"),
    ],
    col_widths=[4.5, 3, 8]
)

doc.add_paragraph()
heading(2, "Frontend")
make_table(
    ["Technologie", "Rôle"],
    [
        ("React 18 + TypeScript", "Interface utilisateur"),
        ("Axios",                 "Appels API REST"),
        ("Recharts",              "Graphiques interactifs (barres, lignes, camemberts)"),
        ("Nginx",                 "Serveur web statique — port 3000"),
    ],
    col_widths=[6, 9.5]
)

doc.add_paragraph()
heading(2, "Infrastructure & Base de données")
make_table(
    ["Composant", "Détail"],
    [
        ("Docker Compose",        "Orchestration de 2 conteneurs : chu_backend, chu_frontend"),
        ("MySQL (XAMPP)",         "Base de données principale — localhost:3306, base chu_ibnsina"),
        ("host.docker.internal",  "Résolution Docker → MySQL hôte Windows"),
        ("SQLite",                "Base de secours utilisée si MySQL indisponible"),
        ("Fichier .env",          "Variables d'environnement : DATABASE_URL, JWT_SECRET, SMTP…"),
    ],
    col_widths=[5, 10.5]
)

page_break()

# ════════════════════════════════════════════════════════════
#  3. ARCHITECTURE
# ════════════════════════════════════════════════════════════
heading(1, "3.  Architecture du Projet")

body("L'application suit une architecture client-serveur découplée :", size=11)
doc.add_paragraph()

code_block(
    "PFE/\n"
    "├── backend/\n"
    "│   ├── main.py              ← API FastAPI (toute la logique métier)\n"
    "│   ├── requirements.txt\n"
    "│   └── Dockerfile\n"
    "├── frontend/\n"
    "│   ├── src/\n"
    "│   │   ├── pages/           ← 15 pages React\n"
    "│   │   ├── components/      ← Composants réutilisables\n"
    "│   │   └── context/         ← ThemeContext, ToastContext\n"
    "│   └── Dockerfile\n"
    "├── data/\n"
    "│   ├── gold/                ← CSV nettoyés (urgences, soins, prédictions)\n"
    "│   └── chu.db               ← SQLite de secours\n"
    "├── models/                  ← Modèles ML sérialisés (.pkl, .json)\n"
    "├── docker-compose.yml\n"
    "└── .env                     ← Secrets et configuration"
)

doc.add_paragraph()
heading(2, "Flux de données")
bullet("Le frontend React envoie des requêtes HTTP à l'API FastAPI via Axios")
bullet("Le backend lit les données depuis MySQL (ou CSV en cache mémoire pour la rapidité)")
bullet("Les modèles ML sont chargés en mémoire au démarrage via joblib")
bullet("Les données temps réel (statuts patients, lits) sont persistées en MySQL")
bullet("L'authentification est gérée par token JWT (expiration 8h par défaut)")

page_break()

# ════════════════════════════════════════════════════════════
#  4. BASE DE DONNÉES
# ════════════════════════════════════════════════════════════
heading(1, "4.  Base de Données")

heading(2, "Configuration de connexion")
code_block("DATABASE_URL = mysql+pymysql://root:@host.docker.internal:3306/chu_ibnsina")

doc.add_paragraph()
heading(2, "Tables principales")
make_table(
    ["Table", "Lignes", "Description"],
    [
        ("urgences",        "~530 904", "Passages aux urgences (2019–2026)"),
        ("soins",           "~737 512", "Actes médicaux et coûts associés"),
        ("lits",            "3 295",    "Inventaire des lits par service et établissement"),
        ("personnel",       "1 086",    "Médecins et personnel soignant"),
        ("etablissements",  "8",        "Sites hospitaliers du réseau CHU"),
        ("patient_statuts", "Dynamique","Suivi temps réel : statut et lit assigné"),
        ("users",           "17",       "Comptes utilisateurs et rôles"),
    ],
    col_widths=[4.5, 3, 8]
)

doc.add_paragraph()
heading(2, "Colonnes clés — table urgences")
make_table(
    ["Colonne", "Type", "Description"],
    [
        ("IPP",                "INT PK",    "Identifiant unique du patient"),
        ("Nom_Complet",        "VARCHAR",   "Nom et prénom"),
        ("CIN",                "VARCHAR",   "Carte d'identité nationale"),
        ("Age",                "INT",       "Âge en années"),
        ("Sexe",               "CHAR(1)",   "M ou F"),
        ("Etablissement",      "VARCHAR",   "Site d'admission"),
        ("Date_Arrivee",       "DATETIME",  "Horodatage d'arrivée"),
        ("Date_Sortie",        "DATETIME",  "Horodatage de sortie"),
        ("Niveau_Triage",      "VARCHAR",   "P1-Critique / P2-Urgent / P3 / P4"),
        ("Motif_Consultation", "VARCHAR",   "Motif médical principal"),
        ("Orientation",        "VARCHAR",   "Hospitalise / Domicile / Transfere / Fugue / Decede"),
        ("Duree_Sejour_min",   "INT",       "Durée totale en minutes"),
        ("Mutuelle",           "VARCHAR",   "AMO, RAMED, Payant, Assurance…"),
        ("Prix_Sejour",        "DECIMAL",   "Coût du séjour (MAD)"),
        ("Saison",             "VARCHAR",   "Hiver / Printemps / Été / Automne"),
    ],
    col_widths=[4.5, 3, 8]
)

page_break()

# ════════════════════════════════════════════════════════════
#  5. PAGES DU TABLEAU DE BORD
# ════════════════════════════════════════════════════════════
heading(1, "5.  Pages du Tableau de Bord (15 pages)")

make_table(
    ["#", "Page", "Fonctionnalité principale"],
    [
        ("1",  "Login",                  "Authentification JWT, gestion des rôles"),
        ("2",  "KPIs Globaux",           "Volume, criticité, durée moyenne, jauge OMS"),
        ("3",  "Analyse Temporelle",     "Séries temporelles, distributions horaire / hebdo / saisonnière"),
        ("4",  "Prédictions IA",         "Prophet 30j, simulateur XGBoost, planification RH/lits par date"),
        ("5",  "Ressources Humaines",    "Flux vs personnel disponible, taux d'occupation, recommandations"),
        ("6",  "Soins Médicaux",         "Actes, coûts, financements (AMO, RAMED, Payant…)"),
        ("7",  "Patients Actuels",       "Liste temps réel 24h, statut, attribution de lits"),
        ("8",  "Admission Patient",      "Formulaire d'admission, calcul automatique des prix"),
        ("9",  "Gestion des Lits",       "Inventaire, disponibilité par service et établissement"),
        ("10", "Gestion Personnel",      "Annuaire médical, spécialités, statuts de service"),
        ("11", "Établissements",         "Fiche descriptive de chaque site hospitalier"),
        ("12", "Carte Géographique",     "Visualisation spatiale des établissements"),
        ("13", "Alertes",                "Seuils configurables : durée, fugue, P1, hospitalisation"),
        ("14", "Tableau de Données",     "Exploration brute, filtres multiples, export Excel/CSV"),
        ("15", "Gestion Utilisateurs",   "CRUD comptes, rôles, permissions par établissement"),
    ],
    col_widths=[1.2, 4.5, 9.8]
)

page_break()

# ════════════════════════════════════════════════════════════
#  6. MODÈLES ML
# ════════════════════════════════════════════════════════════
heading(1, "6.  Modèles Machine Learning")

make_table(
    ["Modèle", "Fichier", "Usage"],
    [
        ("XGBoost",       "xgboost_model.pkl",       "Prédiction du nombre de patients par jour"),
        ("Random Forest", "random_forest_model.pkl",  "Classification du niveau de criticité"),
        ("Prophet (Meta)","prophet_model.json",       "Prévision série temporelle sur 30 jours"),
        ("LabelEncoder",  "label_encoder.pkl",        "Encodage des variables catégorielles"),
    ],
    col_widths=[4, 5.5, 6]
)

doc.add_paragraph()
heading(2, "Pipeline de données")
bullet("Bronze → Silver : nettoyage, correction des valeurs aberrantes, normalisation des dates")
bullet("Silver → Gold  : feature engineering (Saison, Tranche_Horaire, Groupe_Age, Est_Pic…)")
bullet("Gold → Modèles : entraînement, validation croisée, sauvegarde des métriques")

doc.add_paragraph()
heading(2, "Features principales utilisées par XGBoost")
bullet("Jour_Semaine, Mois, Heure_Arrivee, Saison, Jour_Ferie")
bullet("Nb_Medecins_Dispo, Nb_Lits_Dispo")
bullet("Taux_Hospit, Est_Pic (booléen pic de flux)")

page_break()

# ════════════════════════════════════════════════════════════
#  7. FONCTIONNALITÉS CLÉS
# ════════════════════════════════════════════════════════════
heading(1, "7.  Fonctionnalités Clés")

heading(2, "Authentification & Sécurité")
bullet("Tokens JWT avec expiration configurable (défaut : 8 heures)")
bullet("3 rôles : admin (accès total), directeur (filtré par établissement), médecin (lecture seule)")
bullet("Mots de passe hashés SHA-256")
bullet("Chaque directeur ne voit que les données de son établissement")

heading(2, "Admission Patient")
bullet("Formulaire complet : identité, triage, motif, orientation, financement")
bullet("Date d'arrivée = heure locale serveur (protection contre les dates erronées ou futures)")
bullet("Calcul automatique des prix selon la mutuelle et les actes réalisés")
bullet("Établissements chargés dynamiquement depuis la base de données")
bullet("Réinitialisation du formulaire après chaque admission avec conservation de l'établissement")

heading(2, "Patients Actuels (temps réel)")
bullet("Affiche tous les patients admis dans les dernières 24 heures glissantes")
bullet("4 statuts : En triage → En attente → En traitement → Sorti")
bullet("Attribution de lit par sélection dans un menu déroulant (stock depuis table lits)")
bullet("Mise à jour immédiate en base via PATCH /api/patients/{IPP}/statut")

heading(2, "Planification par Date")
bullet("Prévision du nombre de patients selon le jour de la semaine et le mois historique")
bullet("Recommandation du nombre de médecins par spécialité (avec écart vs disponible)")
bullet("Recommandation du nombre de lits par service (proportionnel au taux d'hospitalisation)")

heading(2, "Alertes Configurables")
bullet("4 indicateurs surveillés : durée moyenne séjour, taux de fugue, taux P1, taux hospitalisation")
bullet("Seuils modifiables par l'administrateur en temps réel")
bullet("Niveaux : critique (rouge) / alerte (orange) / ok (vert)")

page_break()

# ════════════════════════════════════════════════════════════
#  8. API
# ════════════════════════════════════════════════════════════
heading(1, "8.  API — Principaux Endpoints")

make_table(
    ["Méthode", "Route", "Description"],
    [
        ("POST",   "/api/login",                        "Authentification, retourne un token JWT"),
        ("GET",    "/api/stats",                        "KPIs globaux (volume, criticité, durée…)"),
        ("GET",    "/api/urgences",                     "Données urgences avec filtres"),
        ("POST",   "/api/patients/admission",           "Admettre un nouveau patient"),
        ("GET",    "/api/patients/aujourd-hui",         "Patients des dernières 24 heures"),
        ("PATCH",  "/api/patients/{IPP}/statut",        "Mettre à jour statut et/ou lit assigné"),
        ("GET",    "/api/predictions",                  "Prévisions Prophet 30 jours"),
        ("GET",    "/api/predictions/planification",    "Planification RH et lits pour une date"),
        ("GET",    "/api/lits",                         "Inventaire complet des lits"),
        ("GET",    "/api/personnel",                    "Liste du personnel médical"),
        ("GET",    "/api/etablissements",               "Liste des établissements"),
        ("GET",    "/api/alertes/check",                "Vérification des indicateurs d'alerte"),
        ("POST",   "/api/alertes/config",               "Modifier les seuils d'alerte"),
        ("GET",    "/api/users",                        "Gestion des utilisateurs (admin)"),
        ("POST",   "/api/users",                        "Créer un compte utilisateur"),
    ],
    col_widths=[2.2, 6.5, 6.8]
)

page_break()

# ════════════════════════════════════════════════════════════
#  9. DÉPLOIEMENT
# ════════════════════════════════════════════════════════════
heading(1, "9.  Déploiement")

heading(2, "Prérequis")
bullet("Docker Desktop (Windows)")
bullet("XAMPP avec MySQL actif sur le port 3306")
bullet("Base de données chu_ibnsina créée dans phpMyAdmin")

heading(2, "Commandes de démarrage")
code_block(
    "# Démarrer l'application complète\n"
    "docker compose up -d --build\n\n"
    "# Reconstruire uniquement le backend\n"
    "docker compose up -d --build backend\n\n"
    "# Consulter les logs\n"
    "docker logs chu_backend --tail=20"
)

doc.add_paragraph()
heading(2, "URLs d'accès")
make_table(
    ["Service", "URL", "Description"],
    [
        ("Frontend",  "http://localhost:3000",       "Interface utilisateur React"),
        ("API REST",  "http://localhost:8000",       "Backend FastAPI"),
        ("Docs API",  "http://localhost:8000/docs",  "Documentation interactive Swagger"),
        ("phpMyAdmin","http://localhost/phpmyadmin", "Interface MySQL (XAMPP)"),
    ],
    col_widths=[3.5, 6, 6]
)

doc.add_paragraph()
heading(2, "Variables d'environnement (.env)")
code_block(
    "DATABASE_URL=mysql+pymysql://root:@host.docker.internal:3306/chu_ibnsina\n"
    "JWT_SECRET=votre_secret_jwt\n"
    "JWT_EXPIRE_MINUTES=480\n"
    "SMTP_HOST=smtp.gmail.com\n"
    "SMTP_PORT=587\n"
    "SMTP_USER=email@gmail.com\n"
    "SMTP_PASSWORD=mot_de_passe_app"
)

page_break()

# ════════════════════════════════════════════════════════════
#  10. PROBLÈMES RÉSOLUS
# ════════════════════════════════════════════════════════════
heading(1, "10.  Problèmes Résolus")

make_table(
    ["Problème", "Cause", "Solution appliquée"],
    [
        (
            "Patients n'apparaissent pas dans 'Patients Actuels'",
            "Filtre strict 'aujourd'hui depuis minuit' + dates erronées stockées (1983, 2024…)",
            "Filtre glissant 24h + clampage automatique : date > 24h ou future → remplacée par l'heure serveur"
        ),
        (
            "Menu déroulant des lits se ferme immédiatement",
            "Composant React défini à l'intérieur de la fonction de rendu → remontage à chaque état",
            "Converti en fonction de rendu renderPatientRow() appelée comme fonction, non comme composant JSX"
        ),
        (
            "Erreur decimal.Decimal * float (Planification IA)",
            "MySQL retourne decimal.Decimal pour COUNT() et SUM() contrairement à SQLite",
            "int(r[1]) sur tous les résultats SQL numériques"
        ),
        (
            "Erreur PRAGMA journal_mode=WAL sur MySQL",
            "PRAGMA est une commande SQLite exécutée sans condition",
            "Guard if IS_SQLITE avant toute commande PRAGMA"
        ),
        (
            "480 000 dates converties en NaT",
            "pandas 3.0 déduit un format unique depuis les premières lignes ; le dataset a deux formats de dates mixtes",
            "Paramètre format='mixed' sur pd.to_datetime()"
        ),
        (
            "Établissement hardcodé dans le formulaire d'admission",
            "Constante ETABLISSEMENTS définie en dur dans le frontend",
            "useEffect + appel /api/etablissements au montage du composant"
        ),
        (
            "Date d'arrivée obsolète si la page reste ouverte",
            "localNow() appelé une seule fois au chargement du module JS",
            "useState(() => ({ ...initialForm, date_arrivee: localNow() })) — initialisation lazy"
        ),
        (
            "ON CONFLICT / INSERT OR IGNORE non supporté MySQL",
            "Syntaxe spécifique à SQLite/PostgreSQL",
            "Branche IS_MYSQL : ON DUPLICATE KEY UPDATE / INSERT IGNORE"
        ),
    ],
    col_widths=[4.5, 5, 6]
)

page_break()

# ════════════════════════════════════════════════════════════
#  11. CONCLUSION
# ════════════════════════════════════════════════════════════
heading(1, "11.  Conclusion")

body(
    "Ce projet représente une solution complète de Business Intelligence et de Data Science "
    "appliquée au domaine hospitalier. Il couvre l'intégralité du cycle de vie d'un projet "
    "de données : collecte, nettoyage ETL, modélisation ML, développement full-stack et "
    "déploiement en conteneurs Docker.",
    size=11
)
doc.add_paragraph()
body(
    "L'application mise en production permet au personnel du CHU Ibn Sina de :",
    size=11
)
bullet("Visualiser les indicateurs clés de performance des urgences en temps réel")
bullet("Prédire la charge journalière avec un horizon de 30 jours")
bullet("Planifier les ressources humaines et matérielles à l'avance")
bullet("Suivre chaque patient de l'admission à la sortie avec attribution de lit")
bullet("Configurer des alertes automatiques sur les indicateurs critiques")

doc.add_paragraph()
body(
    "La migration vers MySQL et l'architecture Docker assurent la robustesse, la portabilité "
    "et la scalabilité de la solution pour un déploiement en environnement hospitalier réel.",
    size=11
)

# ════════════════════════════════════════════════════════════
#  SAUVEGARDE
# ════════════════════════════════════════════════════════════
out = r"c:\Users\Azddine\Desktop\PFE\Rapport_Recapitulatif_PFE_CHU_IbnSina.docx"
doc.save(out)
print(f"Document généré : {out}")
