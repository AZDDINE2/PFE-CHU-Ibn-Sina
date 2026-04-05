"""
Génération du rapport PFE — CHU Ibn Sina Dashboard
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)

# ── Styles helpers ──
def set_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text, color=(26, 59, 219)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=color)
    # underline via border bottom
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '%02X%02X%02X' % color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text, color=(30, 41, 59)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=color)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=11, bold=True, color=(59, 130, 246))
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def info_box(text, bg_rgb=(219, 234, 254)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, size=10, italic=True, color=(30, 58, 138))
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '%02X%02X%02X' % tuple(bg_rgb))
    p._p.get_or_add_pPr().append(shd)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=10, bold=True, color=(255,255,255))
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A3BDB')
        cell._tc.get_or_add_tcPr().append(shd)
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri+1].cells[ci]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=10)
            if ri % 2 == 0:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EFF6FF')
                cell._tc.get_or_add_tcPr().append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ══════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run('UNIVERSITÉ MOHAMMED V DE RABAT')
set_font(run, size=13, bold=True, color=(30, 41, 59))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Faculté des Sciences — Département Informatique')
set_font(run, size=12, color=(100, 116, 139))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1A3BDB')
p._p.get_or_add_pPr().append(shd)
run = p.add_run('  RAPPORT DE PROJET DE FIN D\'ÉTUDES  ')
set_font(run, size=22, bold=True, color=(255,255,255))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after  = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('Tableau de Bord Décisionnel pour la Gestion des Urgences\ndu CHU Ibn Sina de Rabat')
set_font(run, size=18, bold=True, color=(26, 59, 219))
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Réalisé par :')
set_font(run, size=12, bold=True, color=(30, 41, 59))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Azddine')
set_font(run, size=14, bold=True, color=(26, 59, 219))

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Encadrant :  ')
set_font(run, size=11, bold=True)
run2 = p.add_run('Pr. [Nom de l\'encadrant]')
set_font(run2, size=11)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Établissement d\'accueil :  ')
set_font(run, size=11, bold=True)
run2 = p.add_run('CHU Ibn Sina — Rabat, Maroc')
set_font(run2, size=11)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'Année Universitaire 2024–2025')
set_font(run, size=12, bold=True, color=(100, 116, 139))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# RÉSUMÉ
# ══════════════════════════════════════════════════════════════

heading1('Résumé')
body(
    "Ce projet de fin d'études porte sur la conception et le développement d'un tableau de bord "
    "décisionnel pour la gestion des urgences du CHU Ibn Sina de Rabat. Face à l'augmentation "
    "constante du flux de patients aux urgences hospitalières et aux défis organisationnels qui en "
    "découlent, ce travail propose une solution numérique complète intégrant la Business Intelligence, "
    "le Machine Learning et des technologies web modernes."
)
body(
    "La solution développée est une application full-stack composée d'un backend Python (FastAPI) "
    "exposant une API REST, d'une base de données SQLite contenant plus de 100 000 passages aux urgences, "
    "481 fiches de personnel et 2 535 lits répartis sur 9 établissements, et d'un frontend React/TypeScript "
    "offrant une interface moderne et responsive."
)
body(
    "Les fonctionnalités clés incluent : des KPIs en temps réel, une analyse temporelle avancée, "
    "des prédictions ML (Prophet, XGBoost), un système de détection d'anomalies, "
    "une gestion des lits et du personnel, ainsi qu'un module d'intelligence décisionnelle pour "
    "l'orientation des patients. Le système intègre également une gestion des accès multi-rôles "
    "avec filtrage des données par établissement."
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Mots-clés : ')
set_font(run, size=11, bold=True)
run2 = p.add_run('Business Intelligence, Machine Learning, Dashboard, Urgences, CHU, React, FastAPI, Prédiction, Gestion Hospitalière')
set_font(run2, size=11, italic=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════

heading1('Table des Matières')
toc_items = [
    ('1.', 'Introduction', '4'),
    ('  1.1', 'Contexte et problématique', '4'),
    ('  1.2', 'Objectifs du projet', '4'),
    ('  1.3', 'Périmètre', '5'),
    ('2.', 'Analyse de l\'existant', '5'),
    ('  2.1', 'Système actuel de gestion des urgences', '5'),
    ('  2.2', 'Limites identifiées', '5'),
    ('3.', 'Données et Pipeline ETL', '6'),
    ('  3.1', 'Description du dataset', '6'),
    ('  3.2', 'Pipeline de traitement', '6'),
    ('  3.3', 'Statistiques des données', '7'),
    ('4.', 'Architecture Technique', '7'),
    ('  4.1', 'Stack technologique', '7'),
    ('  4.2', 'Architecture globale', '8'),
    ('  4.3', 'API REST — Endpoints', '8'),
    ('5.', 'Fonctionnalités du Dashboard', '9'),
    ('  5.1', 'Authentification et gestion des rôles', '9'),
    ('  5.2', 'KPIs Globaux', '10'),
    ('  5.3', 'Analyse Temporelle', '10'),
    ('  5.4', 'Patients Actuels (Temps Réel)', '10'),
    ('  5.5', 'Vue d\'ensemble RH', '11'),
    ('  5.6', 'Gestion du Personnel', '11'),
    ('  5.7', 'Gestion des Lits', '11'),
    ('  5.8', 'Intelligence Décisionnelle IA', '12'),
    ('6.', 'Modèles de Machine Learning', '13'),
    ('  6.1', 'Prévision d\'affluence (Prophet)', '13'),
    ('  6.2', 'Prédiction de durée (XGBoost)', '13'),
    ('  6.3', 'Évaluation patient et triage', '13'),
    ('  6.4', 'Détection d\'anomalies', '14'),
    ('7.', 'Sécurité et Accès Multi-rôles', '14'),
    ('8.', 'Déploiement (Docker)', '15'),
    ('9.', 'Résultats et Évaluation', '15'),
    ('10.', 'Conclusion et Perspectives', '16'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    dots = '.' * max(1, 60 - len(f'{num}  {title}') - len(page))
    run = p.add_run(f'{num}  {title} {dots} {page}')
    set_font(run, size=11, bold=(not num.startswith(' ')))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 1 — INTRODUCTION
# ══════════════════════════════════════════════════════════════

heading1('1. Introduction')

heading2('1.1 Contexte et problématique')
body(
    "Les services d'urgences hospitalières font face à une pression croissante : augmentation du nombre "
    "de passages, manque de visibilité sur les ressources disponibles, difficulté à anticiper les pics "
    "d'affluence, et absence d'outils décisionnels adaptés. Au Maroc, le CHU Ibn Sina de Rabat, "
    "complexe hospitalier de référence regroupant 9 établissements, traite chaque année des dizaines "
    "de milliers de passages aux urgences dans des conditions qui nécessitent une meilleure coordination "
    "et une aide à la décision en temps réel."
)
body(
    "La gestion de ces urgences repose encore largement sur des outils traditionnels (tableurs, "
    "registres papier) qui ne permettent pas d'exploiter efficacement les données disponibles pour "
    "optimiser les ressources humaines et matérielles."
)

heading2('1.2 Objectifs du projet')
body("Ce projet vise à atteindre les objectifs suivants :")
bullet("Centraliser et visualiser les données des urgences en temps réel")
bullet("Fournir des indicateurs de performance (KPIs) clairs et actionnables")
bullet("Anticiper les flux de patients grâce à des modèles de Machine Learning")
bullet("Optimiser la gestion des lits et du personnel médical")
bullet("Détecter automatiquement les situations anormales")
bullet("Permettre une gestion sécurisée avec des accès différenciés par rôle et par établissement")

heading2('1.3 Périmètre du projet')
body(
    "Le projet couvre l'ensemble du CHU Ibn Sina composé de 9 établissements : CHU Ibn Sina, "
    "Hôpital Al Ayachi, Hôpital Ar-Razi, Hôpital Ibn Sina, Hôpital Moulay Youssef, Hôpital de "
    "Maternité Souissi, Hôpital de Maternité et de Santé Reproductive les Orangers, "
    "Hôpital des Enfants, et Hôpital des Spécialités."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 2 — ANALYSE DE L'EXISTANT
# ══════════════════════════════════════════════════════════════

heading1('2. Analyse de l\'existant')

heading2('2.1 Système actuel de gestion des urgences')
body(
    "Avant ce projet, la gestion des urgences au CHU Ibn Sina reposait sur des processus manuels "
    "et des outils fragmentés. Les données étaient collectées de manière hétérogène dans différents "
    "services sans système de consolidation centralisé. Les rapports de performance étaient produits "
    "manuellement de façon périodique, ne permettant pas de réactivité face aux situations d'urgence."
)

heading2('2.2 Limites identifiées')
bullet("Absence de tableau de bord unifié pour les décideurs")
bullet("Impossibilité de visualiser en temps réel l'état des lits disponibles")
bullet("Aucune prédiction du flux de patients pour anticiper les besoins en ressources")
bullet("Gestion cloisonnée par établissement sans vue consolidée")
bullet("Manque d'alertes automatiques en cas de saturation")
bullet("Aucune traçabilité des actions sur les dossiers patients")

info_box(
    "Solution proposée : Un dashboard web full-stack avec authentification, visualisation en temps réel, "
    "prédictions ML, gestion des ressources et système d'alertes — accessible depuis n'importe quel "
    "navigateur web."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 3 — DONNÉES ET PIPELINE ETL
# ══════════════════════════════════════════════════════════════

heading1('3. Données et Pipeline ETL')

heading2('3.1 Description du dataset')
body(
    "Le dataset principal contient l'historique complet des passages aux urgences du CHU Ibn Sina "
    "sur plusieurs années. Il est enrichi de données sur le personnel médical et la gestion des lits."
)

add_table(
    ['Table', 'Volume', 'Description'],
    [
        ['urgences', '100 103 enregistrements', 'Passages aux urgences avec triage, orientation, durée'],
        ['personnel', '481 membres', 'Staff médical et paramédical par établissement'],
        ['lits',      '2 535 lits', 'Gestion des lits par service et établissement'],
        ['users',     '17 comptes', 'Utilisateurs de la plateforme (multi-rôles)'],
    ],
    col_widths=[3.5, 5, 7]
)

heading2('3.2 Structure des données urgences')
body("Les champs principaux du dataset urgences :")
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['ID_Urgence',        'Numérique',   'Identifiant unique du passage'],
        ['Date_Arrivee',      'DateTime',    'Date et heure d\'arrivée du patient'],
        ['Nom_Complet',       'Texte',       'Identité du patient'],
        ['Age',               'Entier',      'Âge du patient (0-120)'],
        ['Sexe',              'Catégoriel',  'M / F'],
        ['Niveau_Triage',     'Catégoriel',  'P1 Critique → P5 Très peu urgent'],
        ['Motif_Consultation','Texte',       'Raison de la consultation'],
        ['Etablissement',     'Texte',       'Établissement de prise en charge'],
        ['Orientation',       'Catégoriel',  'Hospitalisé / Domicile / Transféré'],
        ['Duree_Sejour_min',  'Numérique',   'Durée totale du séjour en minutes'],
    ],
    col_widths=[4.5, 3, 8]
)

heading2('3.3 Pipeline de traitement des données')
body("Le pipeline ETL transforme les données brutes en données exploitables en 3 étapes :")
bullet("Bronze : Ingestion des données brutes avec détection des anomalies et doublons")
bullet("Silver : Nettoyage, standardisation des formats, traitement des valeurs manquantes")
bullet("Gold : Enrichissement (features ML), agrégations, création des tables de référence")

heading2('3.4 Statistiques clés du dataset')
add_table(
    ['Indicateur', 'Valeur'],
    [
        ['Total passages urgences',       '100 103'],
        ['Établissements couverts',       '9'],
        ['Membres du personnel',          '481'],
        ['Lits total',                    '2 535'],
        ['Lits disponibles',              '1 509 (59.6%)'],
        ['Lits occupés',                  '495 (19.5%)'],
        ['Lits en maintenance',           '531 (20.9%)'],
        ['Période des données',           '2019 – 2024'],
    ],
    col_widths=[8, 7.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 4 — ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════

heading1('4. Architecture Technique')

heading2('4.1 Stack technologique')
add_table(
    ['Couche', 'Technologie', 'Rôle'],
    [
        ['Frontend',    'React 18 + TypeScript',    'Interface utilisateur SPA'],
        ['UI/Charts',   'Recharts + CSS-in-JS',     'Visualisations et composants'],
        ['Backend',     'Python 3.11 + FastAPI',    'API REST et logique métier'],
        ['Base données','SQLite + SQLAlchemy',       'Stockage persistant'],
        ['ML',          'Prophet + XGBoost + scikit-learn', 'Modèles prédictifs'],
        ['Auth',        'JWT (JSON Web Tokens)',     'Authentification sécurisée'],
        ['Déploiement', 'Docker + Docker Compose',  'Conteneurisation'],
        ['Serveur web', 'Nginx',                    'Reverse proxy frontend'],
    ],
    col_widths=[3.5, 5, 7]
)

heading2('4.2 Architecture globale')
body(
    "L'application suit une architecture Client-Serveur découplée. Le frontend React communique "
    "exclusivement avec le backend via des appels HTTP REST. Le backend FastAPI gère l'authentification "
    "JWT, l'accès aux données SQLite, et l'exécution des modèles ML. Docker Compose orchestre "
    "les deux services avec un réseau interne partagé."
)
info_box(
    "Frontend (port 3000 / Nginx) ←→ Backend FastAPI (port 8000) ←→ SQLite DB + Modèles ML\n"
    "Le JWT token inclut le rôle et l'établissement de l'utilisateur pour le filtrage automatique des données."
)

heading2('4.3 API REST — Endpoints principaux')
add_table(
    ['Méthode', 'Endpoint', 'Description'],
    [
        ['POST', '/api/auth/login',              'Authentification, retourne JWT'],
        ['GET',  '/api/kpis/live',               'KPIs en temps réel'],
        ['GET',  '/api/patients/aujourd_hui',     'Patients admis aujourd\'hui'],
        ['GET',  '/api/personnel',               'Liste du personnel (filtré par rôle/établissement)'],
        ['GET',  '/api/lits',                    'Gestion des lits'],
        ['GET',  '/api/lits/stats',              'Statistiques lits par établissement'],
        ['GET',  '/api/personnel/stats',         'Statistiques RH par établissement'],
        ['GET',  '/api/personnel/directeurs',    'Chef de service par établissement'],
        ['GET',  '/api/predictions',             'Prévisions affluence 30 jours'],
        ['GET',  '/api/anomalies',               'Détection d\'anomalies horaires'],
        ['GET',  '/api/lits/recommander',        'Recommandation de lits disponibles'],
        ['POST', '/api/predict/triage',          'Prédiction niveau de triage patient'],
        ['GET',  '/api/stats/comparaison',       'KPIs semaine actuelle vs précédente'],
        ['GET',  '/api/users',                   'Gestion des comptes utilisateurs'],
        ['POST', '/api/users/seed-directors',    'Création automatique des comptes directeurs'],
        ['PATCH','/api/lits/{etab}/{num}',       'Mise à jour statut d\'un lit'],
        ['PATCH','/api/patients/{id}/statut',    'Mise à jour statut patient'],
    ],
    col_widths=[2, 5.5, 8]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 5 — FONCTIONNALITÉS DU DASHBOARD
# ══════════════════════════════════════════════════════════════

heading1('5. Fonctionnalités du Dashboard')

heading2('5.1 Authentification et gestion des rôles')
body(
    "Le système implémente une authentification par JWT avec 7 niveaux d'accès distincts. "
    "Chaque rôle dispose d'un périmètre de pages accessible, et les directeurs d'établissement "
    "voient uniquement les données de leur propre hôpital grâce au filtrage automatique côté backend."
)
add_table(
    ['Rôle', 'Accès', 'Filtrage établissement'],
    [
        ['admin',        'Toutes les pages + gestion utilisateurs', 'Aucun (vision globale)'],
        ['chef_medecin', 'Toutes les pages cliniques',             'Aucun'],
        ['directeur',    'Patients, Personnel, Lits, RH, KPIs',   'Automatique (son établissement)'],
        ['urgentiste',   'Patients actuels, Lits, KPIs',          'Aucun'],
        ['infirmier',    'Patients actuels, Lits',                 'Aucun'],
        ['analyste',     'Analyses, Prédictions, Tableaux',        'Aucun'],
        ['admin_si',     'Utilisateurs, Données techniques',       'Aucun'],
    ],
    col_widths=[3, 7, 5.5]
)

heading2('5.2 KPIs Globaux')
body(
    "La page KPIs Globaux présente les indicateurs de performance en 3 sections :"
)
bullet("KPIs Globaux : total patients, taux P1 (cas critiques), durée moyenne de séjour, taux d'hospitalisation")
bullet("KPIs de la Semaine : comparaison semaine actuelle vs semaine précédente avec badges de tendance (↑ ↓)")
bullet("Distributions : répartition par orientation (pie chart), par niveau de triage (barres), flux horaire")

heading2('5.3 Analyse Temporelle')
body(
    "Visualisation de l'évolution historique du flux de patients avec possibilité d'explorer "
    "les tendances journalières, hebdomadaires, mensuelles et saisonnières. Graphiques interactifs "
    "avec zoom et export des données."
)

heading2('5.4 Patients Actuels (Temps Réel)')
body(
    "Vue en temps réel des patients admis aujourd'hui avec mise à jour automatique toutes les "
    "30 secondes. Fonctionnalités clés :"
)
bullet("Affichage en mode accordéon par établissement ou en tableau global")
bullet("Filtrage par niveau de triage (P1 à P5)")
bullet("Mise à jour du statut patient : En triage / En attente / En traitement / Sorti")
bullet("Assignation d'un lit au patient")
bullet("Compteur d'alertes P1 par établissement")
bullet("Admission de nouveaux patients via formulaire")

heading2('5.5 Vue d\'ensemble RH')
body(
    "Tableau de bord ressources humaines avec accordéon par établissement. Chaque établissement "
    "affiche la fiche du directeur (Chef de service), les statistiques du personnel (médecins, "
    "infirmiers, en service/absent) et les statistiques des lits avec jauge d'occupation."
)

heading2('5.6 Gestion du Personnel')
body(
    "Interface complète de gestion des 481 membres du personnel médical :"
)
bullet("Accordéon par établissement, groupé par catégorie (Médecins / Infirmiers & Aides)")
bullet("KPI cards filtrables par statut (En service, En garde, En congé, Repos)")
bullet("Modification du statut directement depuis le tableau")
bullet("Vue tableau avec filtrage par établissement")

heading2('5.7 Gestion des Lits')
body(
    "Gestion en temps réel des 2 535 lits répartis dans 9 établissements et 5 services :"
)
bullet("Mode accordéon : établissement → service → cartes de lits colorées par statut")
bullet("Statuts : Disponible (vert) / Occupé (rouge) / En maintenance (orange) / Réservé (bleu)")
bullet("Modal d'édition : changement de statut + assignation patient")
bullet("Barre de distribution colorée par établissement")
bullet("KPI cards globaux cliquables pour filtrer par statut")

doc.add_page_break()

heading2('5.8 Intelligence Décisionnelle IA')
body(
    "Page centrale du module d'intelligence artificielle, organisée en 4 onglets :"
)

heading3('Onglet 1 — Prévision Affluence')
body(
    "Prévision du nombre de patients attendus sur les 30 prochains jours basée sur les données "
    "historiques. Affiche : moyenne journalière prévue, pic attendu, nombre de jours critiques, "
    "tendance sur 30 jours, et top 5 des jours à forte affluence."
)

heading3('Onglet 2 — Anomalies du Jour')
body(
    "Détection automatique des anomalies dans le flux horaire. Compare heure par heure "
    "le flux du dernier jour disponible à la moyenne historique. Une anomalie est détectée "
    "si l'écart dépasse ±30%. Les barres rouges signalent les heures anormales."
)

heading3('Onglet 3 — Évaluation Patient')
body(
    "Outil d'aide à la décision clinique : saisie des informations patient (âge, sexe, "
    "heure d'arrivée, antécédents) → prédiction du niveau de triage P1–P5, score de risque, "
    "durée estimée de séjour, et estimation des coûts de prise en charge."
)

heading3('Onglet 4 — Orientation Lits')
body(
    "Recherche en temps réel des lits disponibles par service. L'utilisateur sélectionne "
    "le service requis et obtient instantanément la liste des établissements triés par nombre "
    "de lits disponibles, avec la meilleure recommandation mise en évidence."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 6 — MACHINE LEARNING
# ══════════════════════════════════════════════════════════════

heading1('6. Modèles de Machine Learning')

heading2('6.1 Prévision d\'affluence — Prophet (Facebook)')
body(
    "Le modèle Prophet de Facebook est utilisé pour la prévision du flux de patients. "
    "Il est particulièrement adapté aux séries temporelles avec saisonnalité et jours fériés. "
    "Il génère une prévision sur 30 jours avec intervalles de confiance (yhat_lower / yhat_upper)."
)
bullet("Entrée : série temporelle quotidienne du nombre de passages")
bullet("Sortie : prévision J+1 à J+30 avec borne inférieure et supérieure")
bullet("Avantage : gère naturellement les tendances hebdomadaires et saisonnières")

heading2('6.2 Prédiction de durée de séjour — XGBoost')
body(
    "Un modèle XGBoost (Gradient Boosting) est entraîné pour prédire la durée de séjour "
    "d'un patient aux urgences en minutes. Les features incluent l'âge, le sexe, l'heure "
    "d'arrivée, le niveau de triage, les ressources disponibles et les variables cycliques "
    "(sin/cos de l'heure, du jour, du mois)."
)
bullet("Features : 25 variables dont encodages cycliques et interactions")
bullet("Cible : Duree_Sejour_min")
bullet("Utilisation : estimation de la durée lors de l'évaluation patient")

heading2('6.3 Prédiction de triage — Règles expertes + XGBoost')
body(
    "La prédiction du niveau de triage combine un système de règles expertes (score basé "
    "sur les facteurs de risque cliniques) et le modèle XGBoost pour l'estimation de durée. "
    "Le score de risque est calculé sur 8 points selon :"
)
add_table(
    ['Facteur de risque', 'Points'],
    [
        ['Âge ≥ 70 ans ou ≤ 2 ans',         '+2'],
        ['Âge entre 60 et 69 ans',           '+1'],
        ['Arrivée entre 22h et 6h',           '+1'],
        ['Jour férié',                        '+1'],
        ['Antécédents cardiaques',            '+2'],
        ['Antécédents respiratoires/neuro',   '+2'],
        ['Antécédents diabète/cancer/HTA',    '+1'],
        ['Saison Hiver',                      '+1'],
    ],
    col_widths=[11, 4.5]
)

heading2('6.4 Détection d\'anomalies')
body(
    "Le module de détection d'anomalies compare en temps réel le flux horaire observé "
    "à la moyenne historique calculée sur l'ensemble du dataset. Une heure est considérée "
    "anormale si l'écart relatif dépasse ±30% :"
)
info_box("Écart = (Valeur_Aujourd'hui - Moyenne_Historique) / Moyenne_Historique × 100")
body(
    "Ce système permet de détecter immédiatement les pics d'afflux inhabituels ou les "
    "creux anormaux (incident, grève, etc.) pour une réaction rapide des équipes."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 7 — SÉCURITÉ
# ══════════════════════════════════════════════════════════════

heading1('7. Sécurité et Accès Multi-rôles')

heading2('7.1 Authentification JWT')
body(
    "L'authentification repose sur des JSON Web Tokens (JWT) signés avec algorithme HS256. "
    "Chaque token contient : le nom d'utilisateur (sub), le rôle, l'établissement, "
    "la date d'émission (iat) et d'expiration (exp — 8 heures)."
)
bullet("Les mots de passe sont hashés en SHA-256 avant stockage en base")
bullet("Le token est envoyé dans le header Authorization: Bearer <token>")
bullet("Le backend valide le token à chaque requête protégée")
bullet("L'établissement dans le token permet le filtrage automatique des données pour les directeurs")

heading2('7.2 Comptes directeurs par établissement')
body(
    "Chaque établissement dispose d'un compte directeur (Chef de service) créé automatiquement "
    "via l'endpoint /api/users/seed-directors. Ces comptes sont liés à leur établissement "
    "et ne peuvent voir que les données de celui-ci."
)
add_table(
    ['Compte', 'Établissement'],
    [
        ['d.benali',    'CHU Ibn Sina'],
        ['d.fassi',     'Hôpital Al Ayachi'],
        ['k.lahlou',    'Hôpital Ar-Razi'],
        ['d.tahiri',    'Hôpital Ibn Sina'],
        ['d.lahlou',    'Hôpital Moulay Youssef'],
        ['f.zouheir',   'Hôpital de Maternité Souissi'],
        ['d.bennani',   'Hôpital de Maternité les Orangers'],
        ['d.alaoui',    'Hôpital des Enfants'],
        ['d.benhaddou', 'Hôpital des Spécialités'],
    ],
    col_widths=[5, 10.5]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 8 — DÉPLOIEMENT
# ══════════════════════════════════════════════════════════════

heading1('8. Déploiement (Docker)')

body(
    "L'application est conteneurisée avec Docker et orchestrée via Docker Compose. "
    "Deux conteneurs sont déployés : le backend FastAPI et le frontend React servi par Nginx."
)
add_table(
    ['Conteneur', 'Image', 'Port', 'Rôle'],
    [
        ['chu_backend',  'pfe-backend',  '8000', 'API FastAPI + SQLite + Modèles ML'],
        ['chu_frontend', 'pfe-frontend', '3000', 'React build servi par Nginx'],
    ],
    col_widths=[3.5, 4, 2.5, 5.5]
)

body("Commandes de déploiement :")
p = doc.add_paragraph()
p.style = 'No Spacing'
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), '1E293B')
p._p.get_or_add_pPr().append(shd)
run = p.add_run("  docker compose build && docker compose up -d  ")
set_font(run, name='Courier New', size=10, color=(226, 232, 240))
doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 9 — RÉSULTATS
# ══════════════════════════════════════════════════════════════

heading1('9. Résultats et Évaluation')

heading2('9.1 Couverture fonctionnelle')
add_table(
    ['Module', 'Statut', 'Détail'],
    [
        ['Authentification multi-rôles',    '✓ Complet', '7 rôles, JWT, filtrage établissement'],
        ['KPIs Globaux',                    '✓ Complet', 'Temps réel + comparaison hebdomadaire'],
        ['Analyse Temporelle',              '✓ Complet', 'Historique, distributions, tendances'],
        ['Patients Actuels',                '✓ Complet', 'Temps réel, triage, assignation lits'],
        ['Vue d\'ensemble RH',              '✓ Complet', 'Accordéon par établissement + directeurs'],
        ['Gestion Personnel',               '✓ Complet', '481 agents, statuts, filtres'],
        ['Gestion des Lits',                '✓ Complet', '2535 lits, édition en temps réel'],
        ['Prévisions ML',                   '✓ Complet', 'Prophet 30j, anomalies, triage, orientation'],
        ['Gestion Utilisateurs',            '✓ Complet', 'CRUD, mots de passe, seed directeurs'],
        ['Carte Géographique',              '✓ Présent', 'Localisation des établissements'],
        ['Soins & Coûts',                   '✓ Présent', 'Analyse financière'],
        ['Alertes & Seuils',                '✓ Présent', 'Système d\'alertes configurables'],
    ],
    col_widths=[5.5, 3, 7]
)

heading2('9.2 Performance technique')
bullet("Temps de réponse API : < 200ms pour les endpoints courants")
bullet("100 103 enregistrements interrogeables en temps réel")
bullet("Mise à jour automatique des patients toutes les 30 secondes")
bullet("Interface responsive adaptée aux écrans de 1280px et plus")
bullet("Mode sombre / clair disponible sur toutes les pages")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CHAPITRE 10 — CONCLUSION
# ══════════════════════════════════════════════════════════════

heading1('10. Conclusion et Perspectives')

heading2('10.1 Bilan du projet')
body(
    "Ce projet a permis de développer une solution complète de Business Intelligence et de Machine "
    "Learning pour la gestion des urgences hospitalières du CHU Ibn Sina. L'application couvre "
    "l'intégralité du cycle de vie des données : de l'ingestion et du nettoyage ETL jusqu'à la "
    "visualisation interactive et la prédiction ML, en passant par la gestion opérationnelle "
    "en temps réel."
)
body(
    "La solution répond aux problématiques identifiées : centralisation des données, visibilité "
    "en temps réel sur les ressources, anticipation des flux et sécurisation des accès par rôle "
    "et par établissement."
)

heading2('10.2 Perspectives d\'amélioration')
bullet("Connexion à un système d'information hospitalier (SIH) réel en remplacement du dataset synthétique")
bullet("Entraînement des modèles ML sur données réelles avec validation clinique")
bullet("Application mobile (React Native) pour le personnel soignant")
bullet("Tableau de bord temps réel avec WebSockets pour les mises à jour push")
bullet("Module de gestion des stocks de médicaments et matériel médical")
bullet("Export PDF et Excel des rapports d'activité")
bullet("Intégration avec les systèmes nationaux de santé (RAMED, CNOPS)")

heading2('10.3 Compétences acquises')
bullet("Architecture full-stack : React/TypeScript + FastAPI + SQLite")
bullet("Machine Learning appliqué au domaine médical (Prophet, XGBoost, règles expertes)")
bullet("Business Intelligence : conception de KPIs, visualisation de données")
bullet("Sécurité web : JWT, hachage, contrôle d'accès basé sur les rôles (RBAC)")
bullet("Déploiement Docker / Docker Compose")
bullet("Gestion de projet agile et itérative")

doc.add_page_break()

# ── Bibliographie ──
heading1('Références et Bibliographie')
refs = [
    "[1] FastAPI Documentation — https://fastapi.tiangolo.com",
    "[2] React Documentation — https://react.dev",
    "[3] Prophet: Forecasting at Scale — Facebook Research, 2017",
    "[4] XGBoost: A Scalable Tree Boosting System — Chen & Guestrin, KDD 2016",
    "[5] SQLAlchemy Documentation — https://docs.sqlalchemy.org",
    "[6] JWT RFC 7519 — JSON Web Token Standard",
    "[7] Docker Documentation — https://docs.docker.com",
    "[8] Recharts — A composable charting library — https://recharts.org",
    "[9] OMS — Normes de classification triage hospitalier (Manchester Triage System)",
    "[10] Ministère de la Santé du Maroc — Rapport annuel des urgences hospitalières",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    set_font(run, size=10, color=(30, 41, 59))

# ── Save ──
output_path = r'C:\Users\Azddine\Desktop\PFE\Rapport_PFE_CHU_Ibn_Sina.docx'
doc.save(output_path)
print(f"Rapport généré : {output_path}")
